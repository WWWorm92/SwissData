# clien.py
# SwissTiming Quantum Client (idealized)
# - Connects to Quantum server over TCP (JSON Lines)
# - Shows runs + athletes, supports per-run categories to avoid bib collisions across categories
# - Local roster per category (bib -> name/country) loaded from Excel/CSV and used to override/display names
# - Export selected run / all runs to CSV / JSON snapshot
#
# Protocol:
#   server sends JSONL lines:
#     {"type":"hello","v":1}
#     {"type":"state","state":{...}}  where state is MeetModel._model_to_state()

import argparse
import csv
import datetime
import json
import re
import os
import queue
import socket
import threading
import time
import tkinter as tk
from dataclasses import dataclass, field
from pathlib import Path
from tkinter import filedialog, messagebox
from tkinter import ttk
from typing import Any, Dict, List, Optional, Tuple

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

    # optional docx extras
    WD_ALIGN_PARAGRAPH = None
    WD_ORIENT = None
    WD_TABLE_ALIGNMENT = None
    Cm = None
    Pt = None
    OxmlElement = None
    qn = None


APP_NAME = "Quantum Client"
DEFAULT_CATEGORIES = ["Мужчины", "Женщины", "Юниоры", "Девушки", "Мастерс"]
DIST_PER_SPLIT_M = 125


def now_ts() -> float:
    return time.time()


def safe_int_str(x: Any) -> str:
    s = str(x).strip()
    if not s:
        return ""
    if s.isdigit():
        return s
    num = ""
    for ch in s:
        if ch.isdigit():
            num += ch
        elif num:
            break
    return num


def split_sort_key(x: str):
    x = str(x).strip()
    if x.isdigit():
        return (0, int(x))
    return (1, x)


def fmt_time(sec: Any) -> str:
    if sec is None:
        return ""
    try:
        f = float(sec)
    except Exception:
        return str(sec)
    sign = "-" if f < 0 else ""
    f = abs(f)
    total_ms = int(round(f * 1000))
    s = (total_ms // 1000) % 60
    m = (total_ms // 60000) % 60
    h = total_ms // 3600000
    ms = total_ms % 1000
    if h > 0:
        return f"{sign}{h:d}:{m:02d}:{s:02d}.{ms:03d}"
    if total_ms >= 60000:
        return f"{sign}{m:d}:{s:02d}.{ms:03d}"
    return f"{sign}{total_ms/1000:.3f}"


def fmt_sec_ru(sec: Any) -> str:
    """Seconds with 3 decimals, comma as decimal separator."""
    if sec is None:
        return ""
    try:
        f = float(sec)
    except Exception:
        return str(sec)
    return f"{f:.3f}".replace(".", ",")


def fmt_speed_kmh_ru(distance_m: Optional[int], sec: Any) -> str:
    if not distance_m:
        return ""
    try:
        t = float(sec)
    except Exception:
        return ""
    if t <= 0:
        return ""
    kmh = (float(distance_m) / t) * 3.6
    return f"{kmh:.3f}".replace(".", ",")


_RU_DOW = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]
_RU_MONTH = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def fmt_ru_long_date(dt: datetime.date) -> str:
    # "Воскресенье 8 февраля 2026 г."
    try:
        dow = _RU_DOW[dt.weekday()]
    except Exception:
        dow = ""
    try:
        mon = _RU_MONTH[dt.month - 1]
    except Exception:
        mon = ""
    return f"{dow} {dt.day} {mon} {dt.year} г.".strip()


def _iter_strings(obj: Any, max_depth: int = 4, _depth: int = 0):
    """Yield strings found in nested dict/list structures (best-effort).

    Used for heuristic discipline detection when JSON has inconsistent fields.
    """
    if _depth > max_depth:
        return
    if obj is None:
        return
    if isinstance(obj, str):
        s = obj.strip()
        if s:
            yield s
        return
    if isinstance(obj, (int, float, bool)):
        return
    if isinstance(obj, dict):
        for k, v in obj.items():
            # keys sometimes contain useful hints too
            if isinstance(k, str):
                ks = k.strip()
                if ks:
                    yield ks
            yield from _iter_strings(v, max_depth=max_depth, _depth=_depth + 1)
        return
    if isinstance(obj, (list, tuple)):
        for it in obj:
            yield from _iter_strings(it, max_depth=max_depth, _depth=_depth + 1)
        return


def _git250_score(run: Dict[str, Any]) -> int:
    """Heuristic score for 'Гит 250 м с места' run detection.

    We try (in order): explicit distance fields, discipline/event/name text, and finally inference
    from split structure (standing 250m typically has exactly one intermediate split at 125m).
    """
    score = 0

    # 1) Explicit distance (best signal)
    dm = run_distance_m(run)
    if dm:
        d = abs(int(dm) - 250)
        if d <= 2:
            score += 12
        elif d <= 10:
            score += 8
        elif d <= 30:
            score += 4

    # 2) Text hints (event/discipline/name + nested metadata)
    text = " ".join(s.lower() for s in _iter_strings(run, max_depth=5))
    if re.search(r"(\b250\b|250\s*м|250m)", text):
        score += 4
    if ("гит" in text) or ("time trial" in text) or re.search(r"\btt\b", text):
        score += 4
    if ("с места" in text) or ("standing" in text) or ("from stand" in text) or ("с/м" in text) or ("s/m" in text):
        score += 3
    if ("0.25" in text) and ("km" in text or "км" in text):
        score += 3

    # 3) Split-structure inference (helps when JSON has no distance and no discipline name)
    # Collect numeric split ids across athletes, accepting keys like "1", "S1", "split1".
    split_ids: set[int] = set()
    ath = run.get("athletes") or {}
    if isinstance(ath, dict):
        for a in ath.values():
            if not isinstance(a, dict):
                continue
            splits = a.get("splits") or {}
            if not isinstance(splits, dict):
                continue
            for k in splits.keys():
                m = re.search(r"\d+", str(k))
                if not m:
                    continue
                try:
                    n = int(m.group(0))
                except Exception:
                    continue
                if 0 < n <= 50:
                    split_ids.add(n)

    if split_ids:
        mx = max(split_ids)
        # Typical standing 250: only split 1 (125m) + finish
        if mx == 1:
            score += 6
        elif mx == 2:
            # could be 375/500 depending on increments; lower confidence
            score += 1
        elif mx >= 3:
            score -= 2

        # If explicit distance is absent, infer distance from the maximum split id
        # assuming 125m increments between intermediates.
        if not dm and 1 <= mx <= 10:
            inferred_dm = (mx + 1) * DIST_PER_SPLIT_M
            d2 = abs(int(inferred_dm) - 250)
            if d2 <= 2:
                score += 9
            elif d2 <= 10:
                score += 5

    return score




def _parse_distance_m(val: Any) -> Optional[int]:
    if val is None:
        return None
    if isinstance(val, (int, float)):
        f = float(val)
        if f <= 0:
            return None
        return int(round(f))
    if isinstance(val, dict):
        v = val.get("value")
        if v is None:
            for kk in ("m", "meters", "distance", "dist"):
                if kk in val:
                    v = val.get(kk)
                    break
        unit = str(val.get("unit") or val.get("u") or "").strip().lower()
        dm = _parse_distance_m(v)
        if dm is None:
            return None
        if unit in ("km", "kilometer", "kilometre", "километр", "км"):
            return int(round(dm * 1000))
        return dm

    s = str(val).strip().lower()
    if not s:
        return None
    s = s.replace(",", ".")
    m = re.search(r"(-?\d+(?:\.\d+)?)", s)
    if not m:
        return None
    try:
        num = float(m.group(1))
    except Exception:
        return None
    if num <= 0:
        return None
    if "km" in s or "км" in s:
        num *= 1000
    return int(round(num))


def run_distance_m(run: Dict[str, Any]) -> Optional[int]:
    if not isinstance(run, dict):
        return None
    keys = ("distance_m", "distance", "dist_m", "dist", "meters", "m", "length_m", "len_m")
    for k in keys:
        if k in run:
            dm = _parse_distance_m(run.get(k))
            if dm:
                return dm
    for nested in ("race", "event", "meta", "info"):
        v = run.get(nested)
        if isinstance(v, dict):
            for k in keys:
                if k in v:
                    dm = _parse_distance_m(v.get(k))
                    if dm:
                        return dm
    return None


def athlete_distance_m(ath: Dict[str, Any]) -> Optional[int]:
    if not isinstance(ath, dict):
        return None
    for k in ("distance_m", "dist_m", "distance", "dist", "meters", "m", "length_m", "len_m"):
        if k in ath:
            dm = _parse_distance_m(ath.get(k))
            if dm:
                return dm
    return None


def fmt_dist(meters: Optional[int]) -> str:
    if not meters:
        return ""
    return f"{meters}м"

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)


def app_data_dir() -> Path:
    base = os.environ.get("APPDATA") or str(Path.home())
    return Path(base) / ".quantum_client"


@dataclass
class Settings:
    host: str = "127.0.0.1"
    port: int = 9876
    auto_reconnect: bool = True
    override_server_names: bool = True
    show_distance: bool = True
    categories: List[str] = field(default_factory=lambda: list(DEFAULT_CATEGORIES))
    run_categories: Dict[str, str] = field(default_factory=dict)  # run_key -> category

    secretary_name: str = ""
    chief_judge_name: str = ""
    protocol_mode: str = "time"
    protocol_include_splits: bool = False

    protocol_date: str = ""
    protocol_conditions: str = ""


class SettingsStore:
    def __init__(self):
        self.dir = app_data_dir()
        ensure_dir(self.dir)
        self.path = self.dir / "settings.json"

    def load(self) -> Settings:
        if not self.path.exists():
            return Settings()
        try:
            data = json.loads(self.path.read_text(encoding="utf-8"))
            s = Settings()
            s.host = str(data.get("host", s.host))
            s.port = int(data.get("port", s.port))
            s.auto_reconnect = bool(data.get("auto_reconnect", s.auto_reconnect))
            s.override_server_names = bool(data.get("override_server_names", s.override_server_names))
            s.show_distance = bool(data.get("show_distance", s.show_distance))
            cats = data.get("categories")
            if isinstance(cats, list) and cats:
                s.categories = [str(x) for x in cats if str(x).strip()]
            rc = data.get("run_categories")
            if isinstance(rc, dict):
                s.run_categories = {str(k): str(v) for k, v in rc.items() if str(k).strip()}
            s.secretary_name = str(data.get("secretary_name", s.secretary_name) or "").strip()
            s.chief_judge_name = str(data.get("chief_judge_name", s.chief_judge_name) or "").strip()
            pm = str(data.get("protocol_mode", s.protocol_mode) or "").strip().lower()
            if pm in ("time", "order"):
                s.protocol_mode = pm
            s.protocol_include_splits = bool(data.get("protocol_include_splits", s.protocol_include_splits))

            s.protocol_date = str(data.get("protocol_date", s.protocol_date) or "").strip()
            s.protocol_conditions = str(data.get("protocol_conditions", s.protocol_conditions) or "").strip()
            return s
        except Exception:
            return Settings()

    def save(self, s: Settings):
        try:
            data = {
                "host": s.host,
                "port": s.port,
                "auto_reconnect": s.auto_reconnect,
                "override_server_names": s.override_server_names,
                "show_distance": s.show_distance,
                "categories": s.categories,
                "run_categories": s.run_categories,
                "secretary_name": s.secretary_name,
                "chief_judge_name": s.chief_judge_name,
                "protocol_mode": s.protocol_mode,
                "protocol_include_splits": s.protocol_include_splits,
                "protocol_date": s.protocol_date,
                "protocol_conditions": s.protocol_conditions,
            }
            self.path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass


class RosterStore:
    """
    roster.json structure:
      {
        "categories": {
          "Мужчины": {"97": {"name":"...", "country":"RUS"}, ...},
          ...
        }
      }
    """
    def __init__(self):
        self.dir = app_data_dir()
        ensure_dir(self.dir)
        self.path = self.dir / "roster.json"
        self.data: Dict[str, Dict[str, Dict[str, str]]] = {}
        self.load()

    def load(self):
        self.data = {}
        if not self.path.exists():
            return
        try:
            j = json.loads(self.path.read_text(encoding="utf-8"))
            cats = j.get("categories", {})
            if isinstance(cats, dict):
                for cat, mp in cats.items():
                    if not isinstance(mp, dict):
                        continue
                    cat_s = str(cat).strip()
                    if not cat_s:
                        continue
                    self.data[cat_s] = {}
                    for bib, meta in mp.items():
                        bib_s = safe_int_str(bib)
                        if not bib_s:
                            continue
                        if isinstance(meta, dict):
                            out = {str(k): ("" if v is None else str(v)).strip() for k, v in meta.items()}
                            if "country" in out:
                                out["country"] = out["country"].upper()
                            if "name" not in out:
                                out["name"] = ""
                            if "country" not in out:
                                out["country"] = ""
                            self.data[cat_s][bib_s] = out
                        else:
                            self.data[cat_s][bib_s] = {"name": str(meta).strip(), "country": ""}
        except Exception:
            self.data = {}

    def save(self):
        try:
            out = {"categories": self.data}
            self.path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass

    def ensure_category(self, cat: str):
        cat = str(cat).strip()
        if not cat:
            return
        if cat not in self.data:
            self.data[cat] = {}

    def set_entry(self, cat: str, bib: str, name: str, country: str = "", **extra: str):
        cat = str(cat).strip()
        bib = safe_int_str(bib)
        if not cat or not bib:
            return
        self.ensure_category(cat)
        rec: Dict[str, str] = {"name": (name or "").strip(), "country": (country or "").strip().upper()}
        for k, v in (extra or {}).items():
            ks = str(k).strip()
            if not ks or ks in ("name", "country"):
                continue
            rec[ks] = ("" if v is None else str(v)).strip()
        self.data[cat][bib] = rec

    def get_entry(self, cat: str, bib: str) -> Optional[Dict[str, str]]:
        cat = str(cat).strip()
        bib = safe_int_str(bib)
        if not cat or not bib:
            return None
        return self.data.get(cat, {}).get(bib)

    def delete_entry(self, cat: str, bib: str):
        cat = str(cat).strip()
        bib = safe_int_str(bib)
        if not cat or not bib:
            return
        try:
            del self.data[cat][bib]
        except Exception:
            pass

    def clear_category(self, cat: str):
        cat = str(cat).strip()
        if not cat:
            return
        self.data[cat] = {}


class TcpClientThread(threading.Thread):
    def __init__(self, host: str, port: int, out_queue: "queue.Queue[Dict[str, Any]]", stop_evt: threading.Event):
        super().__init__(daemon=True)
        self.host = host
        self.port = port
        self.q = out_queue
        self.stop_evt = stop_evt
        self.sock: Optional[socket.socket] = None

    def run(self):
        self._emit({"kind": "status", "ok": False, "text": "Подключение..."})
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(5.0)
            s.connect((self.host, self.port))
            s.settimeout(0.8)
            self.sock = s
        except Exception as e:
            self._emit({"kind": "error", "text": f"Не удалось подключиться: {e}"})
            self._emit({"kind": "status", "ok": False, "text": "Отключено"})
            return

        self._emit({"kind": "status", "ok": True, "text": f"TCP {self.host}:{self.port}"})

        buf = b""
        try:
            while not self.stop_evt.is_set():
                try:
                    chunk = self.sock.recv(4096)
                    if not chunk:
                        raise ConnectionError("соединение закрыто")
                    buf += chunk
                except socket.timeout:
                    continue

                while b"\n" in buf:
                    line, buf = buf.split(b"\n", 1)
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        msg = json.loads(line.decode("utf-8", errors="ignore"))
                        self._emit({"kind": "msg", "data": msg, "raw": line.decode("utf-8", errors="ignore")})
                    except Exception:
                        self._emit({"kind": "raw", "text": line.decode("utf-8", errors="ignore")})
        except Exception as e:
            if not self.stop_evt.is_set():
                self._emit({"kind": "error", "text": f"TCP ошибка: {e}"})
        finally:
            try:
                if self.sock:
                    self.sock.close()
            except Exception:
                pass
            self.sock = None
            self._emit({"kind": "status", "ok": False, "text": "Отключено"})

    def _emit(self, item: Dict[str, Any]):
        try:
            self.q.put(item, block=False)
        except Exception:
            pass


class CategoryDialog(tk.Toplevel):
    def __init__(self, master, categories: List[str]):
        super().__init__(master)
        self.title("Категории")
        self.resizable(False, False)
        self.categories = list(categories)
        self.result: Optional[List[str]] = None

        self.lb = tk.Listbox(self, height=10, width=38)
        self.lb.pack(padx=12, pady=(12, 8), fill="both", expand=False)
        for c in self.categories:
            self.lb.insert("end", c)

        row = ttk.Frame(self)
        row.pack(padx=12, pady=(0, 10), fill="x")

        ttk.Button(row, text="Добавить", command=self._add).pack(side="left")
        ttk.Button(row, text="Переименовать", command=self._rename).pack(side="left", padx=(8, 0))
        ttk.Button(row, text="Удалить", command=self._delete).pack(side="left", padx=(8, 0))

        row2 = ttk.Frame(self)
        row2.pack(padx=12, pady=(0, 12), fill="x")
        ttk.Button(row2, text="Отмена", command=self._cancel).pack(side="right")
        ttk.Button(row2, text="OK", command=self._ok).pack(side="right", padx=(0, 8))

        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._cancel)

    def _prompt(self, title: str, initial: str = "") -> Optional[str]:
        d = tk.Toplevel(self)
        d.title(title)
        d.resizable(False, False)
        ttk.Label(d, text=title).pack(padx=12, pady=(12, 6), anchor="w")
        v = tk.StringVar(value=initial)
        e = ttk.Entry(d, textvariable=v, width=38)
        e.pack(padx=12, pady=(0, 10))
        e.focus_set()

        out = {"val": None}

        def ok():
            s = v.get().strip()
            out["val"] = s if s else None
            d.destroy()

        def cancel():
            out["val"] = None
            d.destroy()

        r = ttk.Frame(d)
        r.pack(padx=12, pady=(0, 12), fill="x")
        ttk.Button(r, text="Отмена", command=cancel).pack(side="right")
        ttk.Button(r, text="OK", command=ok).pack(side="right", padx=(0, 8))
        d.grab_set()
        d.wait_window()
        return out["val"]

    def _sel_index(self) -> Optional[int]:
        sel = self.lb.curselection()
        if not sel:
            return None
        return int(sel[0])

    def _add(self):
        s = self._prompt("Новая категория")
        if not s:
            return
        if s in self.categories:
            messagebox.showerror("Ошибка", "Такая категория уже есть")
            return
        self.categories.append(s)
        self.lb.insert("end", s)

    def _rename(self):
        idx = self._sel_index()
        if idx is None:
            return
        cur = self.categories[idx]
        s = self._prompt("Переименовать", initial=cur)
        if not s or s == cur:
            return
        if s in self.categories:
            messagebox.showerror("Ошибка", "Такая категория уже есть")
            return
        self.categories[idx] = s
        self.lb.delete(idx)
        self.lb.insert(idx, s)

    def _delete(self):
        idx = self._sel_index()
        if idx is None:
            return
        cur = self.categories[idx]
        if messagebox.askyesno("Удалить", f"Удалить категорию: {cur}?"):
            self.categories.pop(idx)
            self.lb.delete(idx)

    def _ok(self):
        out = [c for c in self.categories if c.strip()]
        if not out:
            messagebox.showerror("Ошибка", "Нужна хотя бы одна категория")
            return
        self.result = out
        self.destroy()

    def _cancel(self):
        self.result = None
        self.destroy()


class AthleteDialog(tk.Toplevel):
    def __init__(self, master, bib: str, name: str = "", country: str = ""):
        super().__init__(master)
        self.title("Участник")
        self.resizable(False, False)
        self.result: Optional[Dict[str, str]] = None

        bib = str(bib).strip()

        frm = ttk.Frame(self)
        frm.pack(padx=12, pady=12, fill="both", expand=False)

        ttk.Label(frm, text=f"Номер: {bib}").grid(row=0, column=0, columnspan=2, sticky="w")

        ttk.Label(frm, text="Имя").grid(row=1, column=0, sticky="w", pady=(10, 0))
        self.v_name = tk.StringVar(value=str(name or "").strip())
        e_name = ttk.Entry(frm, textvariable=self.v_name, width=42)
        e_name.grid(row=1, column=1, sticky="w", padx=(10, 0), pady=(10, 0))

        ttk.Label(frm, text="Страна/Город").grid(row=2, column=0, sticky="w", pady=(10, 0))
        self.v_country = tk.StringVar(value=str(country or "").strip())
        e_country = ttk.Entry(frm, textvariable=self.v_country, width=18)
        e_country.grid(row=2, column=1, sticky="w", padx=(10, 0), pady=(10, 0))

        btns = ttk.Frame(self)
        btns.pack(padx=12, pady=(0, 12), fill="x")
        ttk.Button(btns, text="Отмена", command=self._cancel).pack(side="right")
        ttk.Button(btns, text="OK", style="Accent.TButton", command=self._ok).pack(side="right", padx=(0, 8))

        self.bind("<Escape>", lambda _e: self._cancel())
        self.bind("<Return>", lambda _e: self._ok())

        self.grab_set()
        e_name.focus_set()
        self.protocol("WM_DELETE_WINDOW", self._cancel)

    def _ok(self):
        name = (self.v_name.get() or "").strip()
        country = (self.v_country.get() or "").strip()
        self.result = {"name": name, "country": country}
        self.destroy()

    def _cancel(self):
        self.result = None
        self.destroy()



class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.geometry("1440x900")
        self.minsize(1150, 720)

        self.settings_store = SettingsStore()
        self.settings = self.settings_store.load()
        self.roster = RosterStore()

        self._q: "queue.Queue[Dict[str, Any]]" = queue.Queue()
        self._stop_evt = threading.Event()
        self._thr: Optional[TcpClientThread] = None
        self._connected = False
        self._last_state_ts = 0.0

        self.state: Dict[str, Any] = {}
        self.selected_run_key: Optional[str] = None

        self._runs_cols: List[str] = []
        self._runs_split_ids: List[str] = []
        self._runs_active_iid: Optional[str] = None
        self._runs_active_col: Optional[str] = None

        self.host_var = tk.StringVar(value=self.settings.host)
        self.port_var = tk.StringVar(value=str(self.settings.port))
        self.status_var = tk.StringVar(value="Отключено")
        self.info_var = tk.StringVar(value="")
        self.auto_reconnect_var = tk.BooleanVar(value=self.settings.auto_reconnect)
        self.override_names_var = tk.BooleanVar(value=self.settings.override_server_names)
        self.show_distance_var = tk.BooleanVar(value=self.settings.show_distance)

        self.run_filter_var = tk.StringVar(value="")
        self.ath_filter_var = tk.StringVar(value="")
        self.category_var = tk.StringVar(value=(self.settings.categories[0] if self.settings.categories else ""))
        self.run_category_var = tk.StringVar(value="")

        self.protocol_type_var = tk.StringVar(value="Произвольно")
        self.protocol_event_var = tk.StringVar(value="")
        self.protocol_round_var = tk.StringVar(value="")
        self.secretary_var = tk.StringVar(value=self.settings.secretary_name)
        self.chief_judge_var = tk.StringVar(value=self.settings.chief_judge_name)
        self.protocol_mode_var = tk.StringVar(value=self.settings.protocol_mode or "time")
        self.protocol_include_splits_var = tk.BooleanVar(value=bool(self.settings.protocol_include_splits))
        self.protocol_notes_var = tk.StringVar(value="")

        self.protocol_date_var = tk.StringVar(value=self.settings.protocol_date)
        self.protocol_conditions_var = tk.StringVar(value=self.settings.protocol_conditions)

        self.protocol_scope_var = tk.StringVar(value="all")  # all | filter | selected

        self._build_style()
        self._build_ui()

        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self.after(80, self._pump)
        self.after(600, self._auto_save_tick)

    def _build_style(self):
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass

        bg = "#0f1117"
        panel = "#151a23"
        panel2 = "#111621"
        fg = "#e6e6e6"
        muted = "#a8b0bf"
        accent = "#3aa0ff"
        ok = "#7ee787"
        danger = "#ff5c5c"
        line = "#242b3a"

        self.colors = {
            "bg": bg,
            "panel": panel,
            "panel2": panel2,
            "fg": fg,
            "muted": muted,
            "accent": accent,
            "ok": ok,
            "danger": danger,
            "line": line,
            "odd": "#121826",
            "even": "#0f1522",
            "head": "#1b2230",
            "select": "#243044",
        }

        self.configure(bg=bg)
        style.configure(".", background=bg, foreground=fg)
        style.configure("TFrame", background=bg)
        style.configure("Card.TFrame", background=panel)
        style.configure("Card2.TFrame", background=panel2)
        style.configure("TLabel", background=bg, foreground=fg)
        style.configure("Muted.TLabel", background=bg, foreground=muted)
        style.configure("BarMuted.TLabel", background=panel, foreground=muted)
        style.configure("BarH2.TLabel", background=panel, foreground=fg, font=("Segoe UI", 13, "bold"))
        style.configure("H1.TLabel", background=bg, foreground=fg, font=("Segoe UI", 16, "bold"))
        style.configure("H2.TLabel", background=bg, foreground=fg, font=("Segoe UI", 13, "bold"))

        style.configure("TButton", background=panel, foreground=fg, borderwidth=0, padding=(12, 9))
        style.map("TButton", background=[("active", self.colors["head"]), ("pressed", self.colors["select"])])

        style.configure("Accent.TButton", background=accent, foreground="#0b0d12")
        style.map("Accent.TButton", background=[("active", "#5bb3ff"), ("pressed", "#2f8fe6")])

        style.configure("TEntry", fieldbackground=panel2, foreground=fg)
        style.configure("TCombobox", fieldbackground=panel2, background=panel2, foreground=fg, arrowcolor=fg)

        style.configure("Treeview", background=panel2, fieldbackground=panel2, foreground=fg, rowheight=34, borderwidth=1, relief="solid")
        style.map("Treeview", background=[("selected", self.colors["select"])], foreground=[("selected", fg)])
        style.configure("Treeview.Heading", background=self.colors["head"], foreground=fg, relief="flat",
                        font=("Segoe UI", 11, "bold"), padding=(10, 10))

    def _build_ui(self):
        c = self.colors
        root = ttk.Frame(self, style="TFrame")
        root.pack(fill="both", expand=True, padx=14, pady=14)

        top = ttk.Frame(root, style="TFrame")
        top.pack(fill="x", pady=(0, 10))
        ttk.Label(top, text="SwissTiming Quantum Client", style="H1.TLabel").pack(side="left")

        bar = ttk.Frame(root, style="Card.TFrame")
        bar.pack(fill="x", pady=(0, 12))
        inner = ttk.Frame(bar, style="Card.TFrame")
        inner.pack(fill="x", padx=12, pady=12)

        ttk.Label(inner, text="Host", style="Muted.TLabel").pack(side="left")
        ttk.Entry(inner, textvariable=self.host_var, width=16).pack(side="left", padx=(6, 12))
        ttk.Label(inner, text="Port", style="Muted.TLabel").pack(side="left")
        ttk.Entry(inner, textvariable=self.port_var, width=7).pack(side="left", padx=(6, 12))

        ttk.Button(inner, text="Подключить", style="Accent.TButton", command=self.connect).pack(side="left", padx=(0, 8))
        ttk.Button(inner, text="Отключить", command=self.disconnect).pack(side="left")

        ttk.Checkbutton(inner, text="Авто-переподключение", variable=self.auto_reconnect_var).pack(side="left", padx=(16, 0))
        ttk.Checkbutton(inner, text="Подменять имена из состава", variable=self.override_names_var).pack(side="left", padx=(16, 0))
        ttk.Checkbutton(inner, text="Показывать дистанцию", variable=self.show_distance_var, command=self._refresh_views).pack(side="left", padx=(16, 0))

        sw = ttk.Frame(inner, style="Card.TFrame")
        sw.pack(side="right")
        self.status_dot = tk.Canvas(sw, width=14, height=14, bg=c["panel"], highlightthickness=0)
        self.status_dot.pack(side="left", padx=(0, 8))
        self._dot_id = self.status_dot.create_oval(2, 2, 12, 12, fill=c["danger"], outline=c["danger"])
        ttk.Label(sw, textvariable=self.status_var, style="BarH2.TLabel").pack(side="left")
        ttk.Label(sw, textvariable=self.info_var, style="BarMuted.TLabel").pack(side="left", padx=(12, 0))

        self.nb = ttk.Notebook(root)
        self.nb.pack(fill="both", expand=True)

        self.tab_runs = ttk.Frame(self.nb, style="TFrame")
        self.tab_roster = ttk.Frame(self.nb, style="TFrame")
        self.tab_export = ttk.Frame(self.nb, style="TFrame")
        self.tab_log = ttk.Frame(self.nb, style="TFrame")
        self.tab_protocol = ttk.Frame(self.nb, style="TFrame")

        self.nb.add(self.tab_runs, text="Заезды")
        self.nb.add(self.tab_roster, text="Состав")
        self.nb.add(self.tab_export, text="Экспорт")
        self.nb.add(self.tab_log, text="Лог")
        self.nb.add(self.tab_protocol, text="Протокол")

        self._build_runs_tab()
        self._build_roster_tab()
        self._build_export_tab()
        self._build_log_tab()
        self._build_protocol_tab()


    def _build_runs_tab(self):
        c = self.colors

        outer = ttk.Frame(self.tab_runs, style="TFrame")
        outer.pack(fill="both", expand=True, padx=12, pady=12)

        top = ttk.Frame(outer, style="Card.TFrame")
        top.pack(fill="x")
        top_in = ttk.Frame(top, style="Card.TFrame")
        top_in.pack(fill="x", padx=12, pady=12)

        ttk.Label(top_in, text="Заезды и участники", style="H2.TLabel").pack(side="left")
        ttk.Button(top_in, text="Категории…", command=self._edit_categories).pack(side="right")

        ctrl = ttk.Frame(outer, style="Card.TFrame")
        ctrl.pack(fill="x", pady=(12, 0))
        ctrl_in = ttk.Frame(ctrl, style="Card.TFrame")
        ctrl_in.pack(fill="x", padx=12, pady=12)

        ttk.Label(ctrl_in, text="Фильтр", style="Muted.TLabel").pack(side="left")
        ttk.Entry(ctrl_in, textvariable=self.run_filter_var).pack(side="left", fill="x", expand=True, padx=(8, 12))
        self.run_filter_var.trace_add("write", lambda *_: (self._render_runs(), self._render_runs_table_text()))

        ttk.Label(ctrl_in, text="Категория заезда", style="Muted.TLabel").pack(side="left")
        self.run_info_var = tk.StringVar(value="—")
        self.run_cat_cb = ttk.Combobox(ctrl_in, textvariable=self.run_category_var, values=self.settings.categories, state="readonly", width=22)
        self.run_cat_cb.pack(side="left", padx=(8, 12))
        self.run_cat_cb.bind("<<ComboboxSelected>>", lambda _e: self._on_run_category_changed())

        ttk.Button(ctrl_in, text="Загрузить состав…", command=self._load_roster_for_current_category).pack(side="right")
        ttk.Button(ctrl_in, text="Экспорт заезда…", command=self._export_selected_run).pack(side="right", padx=(0, 8))

        # Views: table (Excel-like selectable text) + tree
        view_nb = ttk.Notebook(outer)
        view_nb.pack(fill="both", expand=True, pady=(12, 0))

        self.runs_tab_table = ttk.Frame(view_nb, style="TFrame")
        self.runs_tab_tree = ttk.Frame(view_nb, style="TFrame")
        view_nb.add(self.runs_tab_table, text="Таблица")
        view_nb.add(self.runs_tab_tree, text="Дерево")
        self.runs_view_nb = view_nb

        # --- Table view: plain text with tabs (easy to выделять мышкой и копировать) ---
        table_wrap = ttk.Frame(self.runs_tab_table, style="Card.TFrame")
        table_wrap.pack(fill="both", expand=True)
        table_in = ttk.Frame(table_wrap, style="Card.TFrame")
        table_in.pack(fill="both", expand=True, padx=12, pady=12)

        ttk.Label(table_in, text="Табличный вид (как Excel): выделяй мышкой, Ctrl+C копирует.", style="Muted.TLabel").pack(anchor="w", pady=(0, 8))

        t_wrap = ttk.Frame(table_in, style="Card2.TFrame")
        t_wrap.pack(fill="both", expand=True)

        self.runs_text = tk.Text(
            t_wrap,
            wrap="none",
            font=("Consolas", 11),
            bg=c["panel2"],
            fg=c["fg"],
            insertbackground=c["fg"],
            selectbackground=c["select"],
            selectforeground=c["fg"],
            borderwidth=0,
            highlightthickness=0,
        )
        tvsb = ttk.Scrollbar(t_wrap, orient="vertical", command=self.runs_text.yview)
        thsb = ttk.Scrollbar(t_wrap, orient="horizontal", command=self.runs_text.xview)
        self.runs_text.configure(yscrollcommand=tvsb.set, xscrollcommand=thsb.set)

        self.runs_text.pack(side="left", fill="both", expand=True)
        tvsb.pack(side="right", fill="y")
        thsb.pack(side="bottom", fill="x")

        self.runs_text.bind("<Control-c>", self._copy_runs_text_selection)
        self.runs_text.bind("<Control-C>", self._copy_runs_text_selection)
        self.runs_text.bind("<Control-a>", self._runs_text_select_all)
        self.runs_text.bind("<Control-A>", self._runs_text_select_all)
        self.runs_text.bind("<Button-3>", self._runs_text_context_menu)
        self.runs_text.bind("<Button-2>", self._runs_text_context_menu)
        self.runs_text.bind("<Key>", lambda _e: "break")

        # --- Tree view (original) ---
        base_cols = ("cat", "dist", "start", "place", "country", "finish", "status")
        self.runs_tv = ttk.Treeview(self.runs_tab_tree, columns=base_cols, show="tree headings", selectmode="extended")
        self.runs_tv.heading("#0", text="Имя")
        self.runs_tv.heading("cat", text="Кат.")
        self.runs_tv.heading("dist", text="Дист., м")
        self.runs_tv.heading("start", text="Старт")
        self.runs_tv.heading("place", text="Место")
        self.runs_tv.heading("country", text="Стр/Гор")
        self.runs_tv.heading("finish", text="Финиш")
        self.runs_tv.heading("status", text="Статус")

        self.runs_tv.column("#0", width=360, anchor="w")
        self.runs_tv.column("cat", width=120, anchor="center")
        self.runs_tv.column("dist", width=90, anchor="center")
        self.runs_tv.column("start", width=140, anchor="center")
        self.runs_tv.column("place", width=70, anchor="center")
        self.runs_tv.column("country", width=140, anchor="center")
        self.runs_tv.column("finish", width=110, anchor="center")
        self.runs_tv.column("status", width=110, anchor="center")

        self.runs_tv.tag_configure("run_odd", background=c["odd"])
        self.runs_tv.tag_configure("run_even", background=c["even"])
        self.runs_tv.tag_configure("ath", background=c["panel2"])
        self.runs_tv.tag_configure("warn", background="#2a1616")
        self.runs_tv.tag_configure("miss", background="#1b2238")
        self.runs_tv.tag_configure("run_warn", background="#3a2020")

        wrap = ttk.Frame(self.runs_tab_tree, style="Card.TFrame")
        wrap.pack(fill="both", expand=True)

        vsb = ttk.Scrollbar(wrap, orient="vertical", command=self.runs_tv.yview)
        hsb = ttk.Scrollbar(wrap, orient="horizontal", command=self.runs_tv.xview)
        self.runs_tv.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.runs_tv.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")

        self.runs_tv.bind("<<TreeviewSelect>>", self._on_run_select)
        self.runs_tv.bind("<Double-1>", self._on_athlete_double_click)
        self.runs_tv.bind("<Button-1>", self._runs_tree_record_cell)
        self.runs_tv.bind("<Control-c>", self._copy_runs_tree_selection)
        self.runs_tv.bind("<Control-C>", self._copy_runs_tree_selection)
        self.runs_tv.bind("<Control-Shift-c>", lambda _e: (self._copy_runs_selected_runs_full(), "break")[1])
        self.runs_tv.bind("<Control-Shift-C>", lambda _e: (self._copy_runs_selected_runs_full(), "break")[1])
        self.runs_tv.bind("<Button-3>", self._runs_tree_context_menu)
        self.runs_tv.bind("<Button-2>", self._runs_tree_context_menu)

        try:
            view_nb.select(self.runs_tab_table)
        except Exception:
            pass

    def _build_roster_tab(self):
        c = self.colors

        outer = ttk.Frame(self.tab_roster, style="TFrame")
        outer.pack(fill="both", expand=True, padx=12, pady=12)

        top = ttk.Frame(outer, style="Card.TFrame")
        top.pack(fill="x")
        top_in = ttk.Frame(top, style="Card.TFrame")
        top_in.pack(fill="x", padx=12, pady=12)

        ttk.Label(top_in, text="Категория", style="Muted.TLabel").pack(side="left")
        self.roster_cat_cb = ttk.Combobox(top_in, textvariable=self.category_var, values=self.settings.categories, state="readonly", width=22)
        self.roster_cat_cb.pack(side="left", padx=(8, 12))
        self.roster_cat_cb.bind("<<ComboboxSelected>>", lambda _e: self._render_roster())

        ttk.Button(top_in, text="Загрузить Excel/CSV…", style="Accent.TButton", command=self._load_roster_dialog).pack(side="left")
        ttk.Button(top_in, text="Очистить категорию", command=self._clear_roster_category).pack(side="left", padx=(8, 0))
        ttk.Button(top_in, text="Сохранить", command=self._save_roster).pack(side="left", padx=(8, 0))

        ttk.Button(top_in, text="Категории…", command=self._edit_categories).pack(side="right")

        body = ttk.Frame(outer, style="Card.TFrame")
        body.pack(fill="both", expand=True, pady=(12, 0))
        body_in = ttk.Frame(body, style="Card.TFrame")
        body_in.pack(fill="both", expand=True, padx=12, pady=12)

        cols = ("bib", "name", "country")
        self.roster_tv = ttk.Treeview(body_in, columns=cols, show="headings")
        self.roster_tv.heading("bib", text="№")
        self.roster_tv.heading("name", text="Имя")
        self.roster_tv.heading("country", text="Страна/Город")
        self.roster_tv.column("bib", width=90, anchor="center")
        self.roster_tv.column("name", width=520, anchor="w")
        self.roster_tv.column("country", width=180, anchor="center")

        self.roster_tv.tag_configure("odd", background=c["odd"])
        self.roster_tv.tag_configure("even", background=c["even"])

        vsb = ttk.Scrollbar(body_in, orient="vertical", command=self.roster_tv.yview)
        self.roster_tv.configure(yscrollcommand=vsb.set)
        self.roster_tv.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")

        self.roster_tv.bind("<Double-1>", self._on_roster_double_click)

        hint = ttk.Label(outer, text="Двойной клик по строке: редактирование записи в составе.", style="Muted.TLabel")
        hint.pack(anchor="w", pady=(10, 0))

        self._render_roster()

    def _build_export_tab(self):
        outer = ttk.Frame(self.tab_export, style="TFrame")
        outer.pack(fill="both", expand=True, padx=12, pady=12)

        card = ttk.Frame(outer, style="Card.TFrame")
        card.pack(fill="x")
        pad = ttk.Frame(card, style="Card.TFrame")
        pad.pack(fill="x", padx=12, pady=12)

        ttk.Label(pad, text="Экспорт", style="H2.TLabel").pack(anchor="w")

        row = ttk.Frame(pad, style="Card.TFrame")
        row.pack(fill="x", pady=(10, 0))
        ttk.Button(row, text="Экспорт выбранного заезда в CSV…", style="Accent.TButton", command=self._export_selected_run).pack(side="left")
        ttk.Button(row, text="Экспорт всех заездов в CSV…", command=self._export_all_runs).pack(side="left", padx=(8, 0))
        ttk.Button(row, text="Сохранить JSON-снимок…", command=self._save_snapshot_json).pack(side="left", padx=(8, 0))
        ttk.Button(row, text="Открыть JSON-снимок…", command=self._load_snapshot_json).pack(side="left", padx=(8, 0))
        ttk.Button(row, text="Копировать выбранный заезд (текст)", command=self._copy_selected_run_text).pack(side="left", padx=(8, 0))

        self.export_text = tk.Text(outer, height=22, wrap="none",
                                   bg=self.colors["panel2"], fg=self.colors["fg"],
                                   insertbackground=self.colors["fg"],
                                   relief="flat", highlightthickness=1,
                                   highlightbackground=self.colors["line"], highlightcolor=self.colors["accent"])
        self.export_text.pack(fill="both", expand=True, pady=(12, 0))

        self._update_export_preview()
        self._update_protocol_preview()
        self._update_info_bar()
        self._update_protocol_preview()

    def _build_log_tab(self):
        outer = ttk.Frame(self.tab_log, style="TFrame")
        outer.pack(fill="both", expand=True, padx=12, pady=12)

        card = ttk.Frame(outer, style="Card.TFrame")
        card.pack(fill="both", expand=True)
        pad = ttk.Frame(card, style="Card.TFrame")
        pad.pack(fill="both", expand=True, padx=12, pady=12)

        ttk.Label(pad, text="Лог (последние ~1500 строк)", style="H2.TLabel").pack(anchor="w", pady=(0, 10))
        self.log = tk.Text(pad, wrap="none",
                           bg=self.colors["panel2"], fg=self.colors["fg"],
                           insertbackground=self.colors["fg"],
                           relief="flat", highlightthickness=1,
                           highlightbackground=self.colors["line"], highlightcolor=self.colors["accent"])
        self.log.pack(fill="both", expand=True)


    def _build_protocol_tab(self):
        outer = ttk.Frame(self.tab_protocol, style="TFrame")
        outer.pack(fill="both", expand=True, padx=12, pady=12)

        card = ttk.Frame(outer, style="Card.TFrame")
        card.pack(fill="x")
        pad = ttk.Frame(card, style="Card.TFrame")
        pad.pack(fill="x", padx=12, pady=12)

        ttk.Label(pad, text="Протокол (для секретаря)", style="H2.TLabel").grid(row=0, column=0, columnspan=8, sticky="w")

        types = [
            "Произвольно",
            "Гит 250 м с места",
            "Гит 500 м с места",
            "Спринт (квалификация)",
            "Спринт (заезд)",
            "Скретч",
            "Кейрин",
            "Командный спринт",
        ]

        ttk.Label(pad, text="Тип", style="Muted.TLabel").grid(row=1, column=0, sticky="w", pady=(10, 0))
        cb = ttk.Combobox(pad, textvariable=self.protocol_type_var, values=types, state="readonly", width=24)
        cb.grid(row=1, column=1, sticky="w", padx=(8, 14), pady=(10, 0))
        cb.bind("<<ComboboxSelected>>", lambda _e: self._on_protocol_type())

        ttk.Label(pad, text="Дисциплина", style="Muted.TLabel").grid(row=1, column=2, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.protocol_event_var, width=34).grid(row=1, column=3, sticky="w", padx=(8, 14), pady=(10, 0))

        ttk.Label(pad, text="Раунд/заезд", style="Muted.TLabel").grid(row=1, column=4, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.protocol_round_var, width=18).grid(row=1, column=5, sticky="w", padx=(8, 14), pady=(10, 0))

        ttk.Label(pad, text="Дата", style="Muted.TLabel").grid(row=1, column=6, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.protocol_date_var, width=18).grid(row=1, column=7, sticky="w", padx=(8, 0), pady=(10, 0))

        ttk.Label(pad, text="Секретарь", style="Muted.TLabel").grid(row=2, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.secretary_var, width=24).grid(row=2, column=1, sticky="w", padx=(8, 14), pady=(10, 0))

        ttk.Label(pad, text="Главный судья", style="Muted.TLabel").grid(row=2, column=2, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.chief_judge_var, width=34).grid(row=2, column=3, sticky="w", padx=(8, 14), pady=(10, 0))

        ttk.Label(pad, text="Сортировка", style="Muted.TLabel").grid(row=2, column=4, sticky="w", pady=(10, 0))
        frm_mode = ttk.Frame(pad, style="Card.TFrame")
        frm_mode.grid(row=2, column=5, sticky="w", padx=(8, 14), pady=(10, 0))
        ttk.Radiobutton(frm_mode, text="по времени", value="time", variable=self.protocol_mode_var, command=self._update_protocol_preview).pack(side="left")
        ttk.Radiobutton(frm_mode, text="по порядку", value="order", variable=self.protocol_mode_var, command=self._update_protocol_preview).pack(side="left", padx=(10, 0))

        ttk.Checkbutton(pad, text="Отсечки", variable=self.protocol_include_splits_var, command=self._update_protocol_preview).grid(row=2, column=6, sticky="w", pady=(10, 0))

        ttk.Label(pad, text="Условия", style="Muted.TLabel").grid(row=3, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.protocol_conditions_var, width=84).grid(row=3, column=1, columnspan=5, sticky="we", padx=(8, 14), pady=(10, 0))

        ttk.Label(pad, text="Примечание", style="Muted.TLabel").grid(row=4, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(pad, textvariable=self.protocol_notes_var, width=84).grid(row=4, column=1, columnspan=5, sticky="we", padx=(8, 14), pady=(10, 0))

        btns = ttk.Frame(pad, style="Card.TFrame")
        btns.grid(row=4, column=6, columnspan=2, sticky="e", pady=(10, 0))

        ttk.Label(pad, text="Заезды (для протокола)", style="Muted.TLabel").grid(row=5, column=0, sticky="w", pady=(10, 0))
        scope = ttk.Frame(pad, style="Card.TFrame")
        scope.grid(row=5, column=1, columnspan=5, sticky="w", padx=(8, 14), pady=(10, 0))
        ttk.Radiobutton(scope, text="все", value="all", variable=self.protocol_scope_var, command=self._update_protocol_preview).pack(side="left")
        ttk.Radiobutton(scope, text="по фильтру", value="filter", variable=self.protocol_scope_var, command=self._update_protocol_preview).pack(side="left", padx=(10, 0))
        ttk.Radiobutton(scope, text="выделенные в дереве", value="selected", variable=self.protocol_scope_var, command=self._update_protocol_preview).pack(side="left", padx=(10, 0))
        ttk.Button(btns, text="Копировать", style="Accent.TButton", command=self._copy_protocol_text).pack(side="left")
        ttk.Button(btns, text="TXT…", command=self._save_protocol_txt).pack(side="left", padx=(8, 0))
        ttk.Button(btns, text="DOCX…", command=self._save_protocol_docx).pack(side="left", padx=(8, 0))
        ttk.Button(btns, text="Папка…", command=self._save_protocol_docx_folder).pack(side="left", padx=(8, 0))

        for v in (
            self.protocol_event_var,
            self.protocol_round_var,
            self.protocol_date_var,
            self.protocol_conditions_var,
            self.secretary_var,
            self.chief_judge_var,
            self.protocol_notes_var,
        ):
            v.trace_add("write", lambda *_: self._update_protocol_preview())

        self.protocol_text = tk.Text(outer, height=22, wrap="none",
                                     bg=self.colors["panel2"], fg=self.colors["fg"],
                                     insertbackground=self.colors["fg"],
                                     relief="flat", highlightthickness=1,
                                     highlightbackground=self.colors["line"], highlightcolor=self.colors["accent"])
        self.protocol_text.pack(fill="both", expand=True, pady=(12, 0))

        self._update_protocol_preview()

    def _on_protocol_type(self):
        t = (self.protocol_type_var.get() or "").strip()
        if t and t != "Произвольно" and not self.protocol_event_var.get().strip():
            self.protocol_event_var.set(t)
        self._update_protocol_preview()

    def _protocol_sort_rows(self, run_key: str, run: Dict[str, Any]) -> List[Dict[str, Any]]:
        ath = run.get("athletes") or {}
        order = run.get("bib_order")
        if isinstance(order, list) and order:
            bibs = [safe_int_str(b) for b in order if safe_int_str(b)]
        else:
            bibs = [safe_int_str(b) for b in (ath.keys() if isinstance(ath, dict) else [])]
        bibs = [b for b in bibs if b in ath]

        rows = []
        for idx, bib in enumerate(bibs):
            a = ath.get(bib)
            if not isinstance(a, dict):
                continue
            name, country = self._effective_meta(run_key, a)
            finish = a.get("finish")
            status = str(a.get("status") or "").strip()
            fin_f = None
            try:
                fin_f = float(finish) if finish is not None else None
            except Exception:
                fin_f = None
            rows.append({
                "idx": idx,
                "bib": bib,
                "name": name,
                "country": country,
                "finish_raw": finish,
                "finish_f": fin_f,
                "finish": fmt_time(finish),
                "status": status,
                "splits": a.get("splits") or {},
            })

        mode = (self.protocol_mode_var.get() or "time").strip().lower()
        if mode == "order":
            rows.sort(key=lambda r: (r["finish_f"] is None, r["idx"]))
        else:
            rows.sort(key=lambda r: (r["finish_f"] is None, r["finish_f"] if r["finish_f"] is not None else 0.0, r["idx"]))
        return rows

    def _build_protocol_text(self, run_key: str, run: Dict[str, Any]) -> str:
        cat = self.settings.run_categories.get(run_key, "") or self.run_category_var.get().strip() or self.category_var.get().strip()
        start = run.get("start_time") or ""
        ev = str(run.get("event") or run.get("discipline") or run.get("name") or "").strip()
        if not ev:
            ev = self.protocol_event_var.get().strip() or self.protocol_type_var.get().strip()
        rnd = str(run.get("round") or run.get("phase") or run.get("heat") or "").strip()
        if not rnd:
            rnd = self.protocol_round_var.get().strip()
        sec = self.secretary_var.get().strip()
        judge = self.chief_judge_var.get().strip()
        date_line = self.protocol_date_var.get().strip()
        cond = self.protocol_conditions_var.get().strip()
        notes = self.protocol_notes_var.get().strip()

        if not date_line:
            # try derive from run start time
            st = str(run.get("start_time") or run.get("start") or "").strip()
            if st:
                try:
                    s2 = st.replace("Z", "+00:00")
                    dt = datetime.datetime.fromisoformat(s2)
                    date_line = fmt_ru_long_date(dt.date())
                except Exception:
                    pass

        split_ids = self._run_split_ids(run)
        include_splits = bool(self.protocol_include_splits_var.get())

        rows = self._protocol_sort_rows(run_key, run)

        place = 0
        for r in rows:
            st_u = (r.get("status") or "").strip().upper()
            if r.get("finish_f") is not None and st_u not in ("DNS",):
                place += 1
                r["place"] = str(place)
            else:
                r["place"] = ""

        lines = []
        lines.append("ПРОТОКОЛ")
        if ev:
            lines.append(f"Дисциплина: {ev}")
        if date_line:
            lines.append(date_line)
        if cond:
            lines.append(cond)
        dm = run_distance_m(run)
        if dm is not None:
            lines.append(f"Дистанция: {dm} м")
        if rnd:
            lines.append(f"Раунд: {rnd}")
        if cat:
            lines.append(f"Категория: {cat}")
        lines.append(f"Заезд: {run_key}")
        if start:
            lines.append(f"Старт: {start}")
        if judge:
            lines.append(f"Главный судья: {judge}")
        if sec:
            lines.append(f"Секретарь: {sec}")
        lines.append("")

        cols = ["Место", "№", "Имя", "Стр/Гор", "Финиш", "Статус"]
        if include_splits and split_ids:
            cols.append("Отсечки")
        lines.append(";".join(cols))

        for r in rows:
            parts = [r.get("place", ""), r.get("bib", ""), r.get("name", ""), r.get("country", ""), r.get("finish", ""), r.get("status", "")]
            if include_splits and split_ids:
                sp = []
                splits = r.get("splits") if isinstance(r.get("splits"), dict) else {}
                for sid in split_ids:
                    if str(sid) in splits:
                        sp.append(f"S{sid}:{fmt_time(splits.get(str(sid)))}")
                parts.append(" ".join(sp))
            lines.append(";".join(parts))

        if notes:
            lines.append("")
            lines.append(f"Примечание: {notes}")

        return "\n".join(lines) + "\n"

    def _runs_in_display_order(self) -> List[Tuple[str, Dict[str, Any]]]:
        runs = self.state.get("runs", {}) if isinstance(self.state.get("runs"), dict) else {}
        items: List[Tuple[str, Dict[str, Any]]] = []
        for k, v in runs.items():
            if isinstance(v, dict):
                items.append((str(k), v))

        def _parse_start_ts(x: Any) -> Optional[float]:
            if x is None:
                return None
            if isinstance(x, (int, float)):
                return float(x)
            s = str(x).strip()
            if not s:
                return None
            try:
                # supports "YYYY-MM-DD HH:MM:SS" and ISO formats
                s2 = s.replace("Z", "+00:00")
                return datetime.datetime.fromisoformat(s2).timestamp()
            except Exception:
                return None

        def _sort_key(item: Tuple[str, Dict[str, Any]]):
            k, run = item

            for fld in ("order", "seq", "index", "heat", "run_no", "number"):
                if fld in run:
                    try:
                        return (0, int(run.get(fld)))
                    except Exception:
                        pass

            ts = _parse_start_ts(run.get("start_ts") or run.get("start_time") or run.get("start") or run.get("ts"))
            if ts is not None:
                return (1, ts)

            num = safe_int_str(k)
            if num:
                try:
                    return (2, int(num))
                except Exception:
                    pass

            return (3, k)

        items.sort(key=_sort_key)
        return items

    def _get_protocol_items(self):
        """
        Базовый список заездов для протоколов.
        Возвращает список (run_key, run_dict).
        """
        st = self.state if isinstance(getattr(self, "state", None), dict) else {}
        runs = st.get("runs", {})

        # на всякий случай, если state вложен
        if not isinstance(runs, dict) and isinstance(st.get("state"), dict):
            runs = st["state"].get("runs", {})

        if not isinstance(runs, dict):
            return []

        # порядок как в JSON (Python 3.7+ сохраняет порядок вставки)
        return [(rk, runs.get(rk) or {}) for rk in runs.keys()]

    def _get_protocol_items_scoped(self) -> List[Tuple[str, Dict[str, Any]]]:
        items = self._get_protocol_items()
        scope = (self.protocol_scope_var.get() or "all").strip().lower()
        if scope == "filter":
            flt = self.run_filter_var.get().strip().lower()
            if not flt:
                return items
            out: List[Tuple[str, Dict[str, Any]]] = []
            for run_key, run in items:
                start = str(run.get("start_time") or "").strip()
                ev = str(run.get("event") or run.get("discipline") or run.get("name") or "").strip()
                rnd = str(run.get("round") or run.get("phase") or run.get("heat") or "").strip()
                cat = self.settings.run_categories.get(run_key, "") or ""
                dm = run_distance_m(run)
                dm_s = str(dm) if dm is not None else ""
                run_hay = f"{run_key} {ev} {rnd} {cat} {start} {dm_s}".lower()
                if flt in run_hay:
                    out.append((run_key, run))
                    continue
                ath = run.get("athletes") or {}
                matched = False
                if isinstance(ath, dict):
                    for bib, a in ath.items():
                        if not isinstance(a, dict):
                            continue
                        name, country = self._effective_meta(run_key, a)
                        finish = fmt_time(a.get("finish"))
                        status = str(a.get("status") or "")
                        hay = f"{bib} {name} {country} {finish} {status}".lower()
                        if flt in hay:
                            matched = True
                            break
                if matched:
                    out.append((run_key, run))
            return out
        if scope == "selected":
            keys = []
            try:
                keys = self._runs_selected_run_keys()
            except Exception:
                keys = []
            if not keys:
                return []
            want = set(keys)
            return [(rk, rr) for rk, rr in items if rk in want]
        return items




    def _build_all_protocol_text(self) -> str:
        items = self._get_protocol_items_scoped()
        if not items:
            return "Нет данных.\n"
        blocks: List[str] = []
        for run_key, run in items:
            blocks.append(self._build_protocol_text(run_key, run).rstrip())
        sep = "\n" + ("-" * 72) + "\n"
        return (sep.join(blocks) + "\n") if blocks else "Нет данных.\n"

    def _update_protocol_preview(self):
        if not hasattr(self, "protocol_text"):
            return
        txt = self._build_all_protocol_text()
        try:
            cur = self.protocol_text.get("1.0", "end")
            if cur != txt:
                self.protocol_text.delete("1.0", "end")
                self.protocol_text.insert("1.0", txt)
        except Exception:
            pass


    def _copy_protocol_text(self):
        txt = self._build_all_protocol_text()
        if not txt.strip():
            return
        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
        except Exception:
            pass
        try:
            self.protocol_text.delete("1.0", "end")
            self.protocol_text.insert("1.0", txt)
        except Exception:
            pass


    def _save_protocol_txt(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text", "*.txt")],
                                            initialfile="protocol_all.txt")
        if not path:
            return
        txt = self._build_all_protocol_text()
        Path(path).write_text(txt, encoding="utf-8")
        messagebox.showinfo("Готово", f"Сохранено: {path}")

    def _docx_setup_page(self, doc):
        """Best-effort page setup for wide protocols."""
        if WD_ORIENT is None or Cm is None:
            return
        try:
            section = doc.sections[-1]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
        except Exception:
            pass

    def _docx_set_default_font(self, doc, font_name: str = "Times New Roman", size_pt: int = 10):
        if Pt is None:
            return
        try:
            st = doc.styles["Normal"]
            st.font.name = font_name
            st.font.size = Pt(size_pt)
            # Ensure Cyrillic font mapping (Word can ignore .font.name otherwise)
            if qn is not None:
                rpr = st._element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rfonts.set(qn(attr), font_name)
        except Exception:
            pass

    def _docx_set_paragraph_spacing_0(self, p):
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        except Exception:
            pass

    def _docx_set_cell_shading(self, cell, fill: str = "D9D9D9"):
        if OxmlElement is None or qn is None:
            return
        try:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)
        except Exception:
            pass

    def _docx_set_table_borders(self, table, outer_sz: int = 12, inner_sz: int = 6, color: str = "000000"):
        """Sizes are in 1/8 pt (Word units)."""
        if OxmlElement is None or qn is None:
            return
        try:
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                tbl.insert(0, tbl_pr)
            tbl_borders = tbl_pr.find(qn("w:tblBorders"))
            if tbl_borders is None:
                tbl_borders = OxmlElement("w:tblBorders")
                tbl_pr.append(tbl_borders)

            def _edge(tag: str, sz: int):
                el = tbl_borders.find(qn(f"w:{tag}"))
                if el is None:
                    el = OxmlElement(f"w:{tag}")
                    tbl_borders.append(el)
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), str(int(sz)))
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)

            for tag in ("top", "left", "bottom", "right"):
                _edge(tag, outer_sz)
            _edge("insideH", inner_sz)
            _edge("insideV", inner_sz)
        except Exception:
            pass

    def _docx_set_cell(self, cell, text: str, align: str = "center", size_pt: int = 8, bold: bool = False):
        try:
            cell.text = ""
            p = cell.paragraphs[0]
            self._docx_set_paragraph_spacing_0(p)
            if WD_ALIGN_PARAGRAPH is not None:
                if align == "left":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text or ""))
            if Pt is not None:
                r.font.size = Pt(size_pt)
            r.bold = bool(bold)
        except Exception:
            try:
                cell.text = str(text or "")
            except Exception:
                pass

    def _docx_add_center_line(self, doc, text: str, size_pt: int = 14, bold: bool = True):
        p = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._docx_set_paragraph_spacing_0(p)
        run = p.add_run(str(text or ""))
        run.bold = bool(bold)
        if Pt is not None:
            run.font.size = Pt(size_pt)
        return p

    def _docx_git250_page(self, doc, run_key: str, run: Dict[str, Any], ev: str, stage: str,
                          date_line: str, conditions: str,
                          sec: str, judge: str, notes: str):
        self._docx_setup_page(doc)
        self._docx_set_default_font(doc, font_name="Times New Roman", size_pt=10)

        # derive date if not provided
        if not date_line:
            st = str(run.get("start_time") or run.get("start") or "").strip()
            if st:
                try:
                    s2 = st.replace("Z", "+00:00")
                    dt = datetime.datetime.fromisoformat(s2)
                    date_line = fmt_ru_long_date(dt.date())
                except Exception:
                    pass

        title = (ev or "ГИТ 250 м с/м").strip()
        if "гит" in title.lower() and "250" in title:
            title = "ГИТ 250 м с/м"
        self._docx_add_center_line(doc, title.upper(), size_pt=14, bold=True)
        if stage:
            self._docx_add_center_line(doc, stage, size_pt=12, bold=True)
        self._docx_add_center_line(doc, "РЕЗУЛЬТАТЫ", size_pt=12, bold=True)
        if date_line:
            self._docx_add_center_line(doc, date_line, size_pt=11, bold=False)
        if conditions:
            self._docx_add_center_line(doc, conditions, size_pt=10, bold=False)

        # small spacer (as in printed Excel-like protocol)
        try:
            p = doc.add_paragraph("")
            self._docx_set_paragraph_spacing_0(p)
        except Exception:
            pass

        ath = run.get("athletes") or {}
        order = run.get("bib_order")
        if isinstance(order, list) and order:
            bibs = [safe_int_str(b) for b in order if safe_int_str(b)]
        else:
            bibs = [safe_int_str(b) for b in (ath.keys() if isinstance(ath, dict) else [])]
        bibs = [b for b in bibs if b and b in ath]

        split_ids = self._run_split_ids(run)
        first_sid = split_ids[0] if split_ids else "1"

        rows = []
        for idx, bib in enumerate(bibs):
            a = ath.get(bib)
            if not isinstance(a, dict):
                continue
            meta = self._effective_meta_full(run_key, a)
            finish = a.get("finish")
            status = str(a.get("status") or "").strip()
            fin_f = None
            try:
                fin_f = float(finish) if finish is not None else None
            except Exception:
                fin_f = None

            splits = a.get("splits")
            s1_f = None
            try:
                if isinstance(splits, dict):
                    v = splits.get(str(first_sid))
                    s1_f = float(v) if v is not None else None
                elif isinstance(splits, (list, tuple)):
                    if str(first_sid).isdigit():
                        j = int(str(first_sid)) - 1
                        if 0 <= j < len(splits):
                            s1_f = float(splits[j])
            except Exception:
                s1_f = None

            seg2 = None
            if fin_f is not None and s1_f is not None:
                seg2 = fin_f - s1_f

            rows.append({
                "idx": idx,
                "bib": bib,
                "meta": meta,
                "finish_f": fin_f,
                "finish": fmt_sec_ru(fin_f) if fin_f is not None else "",
                "status": status,
                "s1": fmt_sec_ru(s1_f) if s1_f is not None else "",
                "s2": fmt_sec_ru(seg2) if seg2 is not None else "",
                "speed": fmt_speed_kmh_ru(250, fin_f) if fin_f is not None else "",
            })

        # sort by time
        rows.sort(key=lambda r: (r["finish_f"] is None, r["finish_f"] if r["finish_f"] is not None else 0.0, r["idx"]))

        place = 0
        for r in rows:
            st_u = (r.get("status") or "").strip().upper()
            if r.get("finish_f") is not None and st_u not in ("DNS",):
                place += 1
                r["place"] = str(place)
            else:
                r["place"] = ""

        cols = [
            "Ме\nст.",
            "№",
            "Фамилия Имя Отчество",
            "Организация",
            "Дата\nРождения",
            "Разряд",
            "Регион",
            "125 м",
            "125-\n125м",
            "Результ\nат",
            "Ср.ск.\nт-н",
        ]
        table = doc.add_table(rows=1, cols=len(cols))
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        try:
            table.autofit = False
        except Exception:
            pass

        if WD_TABLE_ALIGNMENT is not None:
            try:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            except Exception:
                pass
        self._docx_set_table_borders(table, outer_sz=12, inner_sz=6)

        widths_cm = [0.9, 0.9, 5.8, 5.6, 2.3, 1.4, 2.6, 1.6, 1.9, 1.8, 1.8]
        if Cm is not None:
            for j, w in enumerate(widths_cm):
                for cell in table.columns[j].cells:
                    cell.width = Cm(w)

        hdr = table.rows[0].cells
        for j, c in enumerate(cols):
            self._docx_set_cell_shading(hdr[j], fill="D9D9D9")
            self._docx_set_cell(hdr[j], c, align="center", size_pt=8, bold=True)

        for r in rows:
            m = r.get("meta") or {}
            line = [
                r.get("place", ""),
                r.get("bib", ""),
                m.get("name", ""),
                m.get("org", ""),
                m.get("dob", ""),
                m.get("rank", ""),
                (m.get("region") or m.get("country") or ""),
                r.get("s1", ""),
                r.get("s2", ""),
                r.get("finish", "") if r.get("finish") else (r.get("status") or ""),
                r.get("speed", ""),
            ]
            cells = table.add_row().cells
            for j, txt in enumerate(line):
                align = "left" if j in (2, 3) else "center"
                self._docx_set_cell(cells[j], txt, align=align, size_pt=8, bold=False)

        # Signature block (optional, stays at bottom in printouts)
        if judge or sec or notes:
            try:
                doc.add_paragraph("")
            except Exception:
                pass
        if notes:
            try:
                p = doc.add_paragraph(notes)
                self._docx_set_paragraph_spacing_0(p)
            except Exception:
                pass
        if judge or sec:
            try:
                sig = doc.add_table(rows=1, cols=2)
                try:
                    sig.autofit = False
                except Exception:
                    pass
                if Cm is not None:
                    for cell in sig.columns[0].cells:
                        cell.width = Cm(13.0)
                    for cell in sig.columns[1].cells:
                        cell.width = Cm(13.0)
                self._docx_set_table_borders(sig, outer_sz=0, inner_sz=0)
                left = sig.rows[0].cells[0]
                right = sig.rows[0].cells[1]
                self._docx_set_cell(left, f"Главный судья: {judge}" if judge else "", align="left", size_pt=10, bold=False)
                self._docx_set_cell(right, f"Секретарь: {sec}" if sec else "", align="left", size_pt=10, bold=False)
            except Exception:
                pass


    def _save_protocol_docx(self):
        if Document is None:
            messagebox.showerror("Ошибка", "python-docx не установлен")
            return

        items = self._get_protocol_items_scoped()
        if not items:
            messagebox.showerror("Ошибка", "Нет данных по заездам")
            return

        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")],
                                            initialfile="protocol_all.docx")
        if not path:
            return

        sec = self.secretary_var.get().strip()
        judge = self.chief_judge_var.get().strip()
        notes = self.protocol_notes_var.get().strip()
        proto_type = (self.protocol_type_var.get() or "").strip()
        date_line_global = (self.protocol_date_var.get() or "").strip()
        conditions_global = (self.protocol_conditions_var.get() or "").strip()

        if proto_type == "Гит 250 м с места":
            # If the user limited the scope (filter/selected), trust that selection.
            scope = (self.protocol_scope_var.get() or "all").strip().lower()
            if scope in ("filter", "selected") and items:
                pass
            else:
                scored: List[Tuple[int, str, Dict[str, Any]]] = []
                for rk, rr in items:
                    scored.append((_git250_score(rr), rk, rr))

                best = max((s for s, _, _ in scored), default=0)
                # thresholds: prefer strong distance/split inference matches, then keyword matches
                if any(s >= 12 for s, _, _ in scored):
                    picked = [(rk, rr) for s, rk, rr in scored if s >= 12]
                elif any(s >= 9 for s, _, _ in scored):
                    picked = [(rk, rr) for s, rk, rr in scored if s >= 9]
                elif any(s >= 5 for s, _, _ in scored):
                    picked = [(rk, rr) for s, rk, rr in scored if s >= 5]
                elif best > 0:
                    picked = [(rk, rr) for s, rk, rr in scored if s == best]
                else:
                    picked = []

                if picked:
                    items = picked
                else:
                    # do not block the user: fall back to the first run, but warn
                    messagebox.showwarning(
                        "Предупреждение",
                        "Не удалось автоматически определить заезды 'Гит 250 м' по JSON.\n"
                        "Сформирую протокол по первому доступному заезду.\n"
                        "Рекомендация: добавь в JSON поле distance=250 или event/discipline с 'Гит 250'.",
                    )
                    items = [items[0]]

        doc = Document()

        for i, (run_key, run) in enumerate(items):
            if i > 0:
                doc.add_page_break()

            cat = self.settings.run_categories.get(run_key, "") or self.run_category_var.get().strip() or self.category_var.get().strip()
            start = run.get("start_time") or ""

            ev = str(run.get("event") or run.get("discipline") or run.get("name") or "").strip()
            if not ev:
                ev = self.protocol_event_var.get().strip() or self.protocol_type_var.get().strip()

            rnd = str(run.get("round") or run.get("phase") or run.get("heat") or "").strip()
            if not rnd:
                rnd = self.protocol_round_var.get().strip()

            dm = run_distance_m(run)

            if proto_type == "Гит 250 м с места":
                ev_title = self.protocol_event_var.get().strip() or "ГИТ 250 м с/м"
                stage = self.protocol_round_var.get().strip() or rnd
                self._docx_git250_page(
                    doc,
                    run_key=run_key,
                    run=run,
                    ev=ev_title,
                    stage=stage,
                    date_line=date_line_global,
                    conditions=conditions_global,
                    sec=sec,
                    judge=judge,
                    notes=notes,
                )
                continue

            split_ids = self._run_split_ids(run)
            include_splits = bool(self.protocol_include_splits_var.get())

            rows = self._protocol_sort_rows(run_key, run)
            place = 0
            for r in rows:
                st_u = (r.get("status") or "").strip().upper()
                if r.get("finish_f") is not None and st_u not in ("DNS",):
                    place += 1
                    r["place"] = str(place)
                else:
                    r["place"] = ""

            doc.add_heading("Протокол", level=1)
            if ev:
                doc.add_paragraph(f"Дисциплина: {ev}")

            date_line = date_line_global
            if not date_line:
                st = str(run.get("start_time") or run.get("start") or "").strip()
                if st:
                    try:
                        s2 = st.replace("Z", "+00:00")
                        dt = datetime.datetime.fromisoformat(s2)
                        date_line = fmt_ru_long_date(dt.date())
                    except Exception:
                        pass
            if date_line:
                doc.add_paragraph(date_line)
            if conditions_global:
                doc.add_paragraph(conditions_global)
            if dm is not None:
                doc.add_paragraph(f"Дистанция: {dm} м")
            if rnd:
                doc.add_paragraph(f"Раунд: {rnd}")
            if cat:
                doc.add_paragraph(f"Категория: {cat}")
            doc.add_paragraph(f"Заезд: {run_key}")
            if start:
                doc.add_paragraph(f"Старт: {start}")
            if judge:
                doc.add_paragraph(f"Главный судья: {judge}")
            if sec:
                doc.add_paragraph(f"Секретарь: {sec}")
            if notes:
                doc.add_paragraph(f"Примечание: {notes}")

            cols = ["Место", "№", "Имя", "Стр/Гор", "Финиш", "Статус"]
            if include_splits and split_ids:
                cols.append("Отсечки")

            table = doc.add_table(rows=1, cols=len(cols))
            hdr = table.rows[0].cells
            for j, c in enumerate(cols):
                hdr[j].text = c

            for r in rows:
                cells = table.add_row().cells
                cells[0].text = r.get("place", "")
                cells[1].text = r.get("bib", "")
                cells[2].text = r.get("name", "")
                cells[3].text = r.get("country", "")
                cells[4].text = r.get("finish", "")
                cells[5].text = r.get("status", "")
                if include_splits and split_ids:
                    sp = []
                    splits = r.get("splits") if isinstance(r.get("splits"), dict) else {}
                    for sid in split_ids:
                        if str(sid) in splits:
                            sp.append(f"S{sid}:{fmt_time(splits.get(str(sid)))}")
                    cells[-1].text = " ".join(sp)

        doc.save(path)
        messagebox.showinfo("Готово", f"Сохранено: {path}")

    def _save_protocol_docx_folder(self):
        if Document is None:
            messagebox.showerror("Ошибка", "python-docx не установлен")
            return

        items = self._get_protocol_items_scoped()
        if not items:
            messagebox.showerror("Ошибка", "Нет данных по заездам")
            return

        folder = filedialog.askdirectory(title="Выберите папку для сохранения протоколов")
        if not folder:
            return

        sec = self.secretary_var.get().strip()
        judge = self.chief_judge_var.get().strip()
        notes = self.protocol_notes_var.get().strip()
        proto_type = (self.protocol_type_var.get() or "").strip()
        date_line_global = (self.protocol_date_var.get() or "").strip()
        conditions_global = (self.protocol_conditions_var.get() or "").strip()

        # If generating Git250 and scope is 'all', try to auto-pick matching runs (same logic as in DOCX…).
        if proto_type == "Гит 250 м с места":
            scope = (self.protocol_scope_var.get() or "all").strip().lower()
            if scope not in ("filter", "selected"):
                scored: List[Tuple[int, str, Dict[str, Any]]] = []
                for rk, rr in items:
                    scored.append((_git250_score(rr), rk, rr))
                best = max((s for s, _, _ in scored), default=0)
                if any(s >= 12 for s, _, _ in scored):
                    items = [(rk, rr) for s, rk, rr in scored if s >= 12]
                elif any(s >= 9 for s, _, _ in scored):
                    items = [(rk, rr) for s, rk, rr in scored if s >= 9]
                elif any(s >= 5 for s, _, _ in scored):
                    items = [(rk, rr) for s, rk, rr in scored if s >= 5]
                elif best > 0:
                    items = [(rk, rr) for s, rk, rr in scored if s == best]
                else:
                    messagebox.showwarning(
                        "Предупреждение",
                        "Не удалось автоматически определить заезды 'Гит 250 м' по JSON.\n"
                        "Сохраню протокол по первому доступному заезду.\n"
                        "Рекомендация: добавь в JSON поле distance=250 или event/discipline с 'Гит 250'.",
                    )
                    items = [items[0]]

        def _safe_name(s: str) -> str:
            s = (s or "").strip()
            if not s:
                return "protocol"
            s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
            s = re.sub(r"\s+", " ", s).strip()
            return s[:120] if len(s) > 120 else s

        saved = 0
        for run_key, run in items:
            try:
                doc = Document()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось создать DOCX: {e}")
                return

            cat = self.settings.run_categories.get(run_key, "") or self.run_category_var.get().strip() or self.category_var.get().strip()
            start = run.get("start_time") or ""

            ev = str(run.get("event") or run.get("discipline") or run.get("name") or "").strip()
            if not ev:
                ev = self.protocol_event_var.get().strip() or self.protocol_type_var.get().strip()

            rnd = str(run.get("round") or run.get("phase") or run.get("heat") or "").strip()
            if not rnd:
                rnd = self.protocol_round_var.get().strip()

            dm = run_distance_m(run)

            if proto_type == "Гит 250 м с места":
                ev_title = self.protocol_event_var.get().strip() or "ГИТ 250 м с/м"
                stage = self.protocol_round_var.get().strip() or rnd
                self._docx_git250_page(
                    doc,
                    run_key=run_key,
                    run=run,
                    ev=ev_title,
                    stage=stage,
                    date_line=date_line_global,
                    conditions=conditions_global,
                    sec=sec,
                    judge=judge,
                    notes=notes,
                )
            else:
                split_ids = self._run_split_ids(run)
                include_splits = bool(self.protocol_include_splits_var.get())

                rows = self._protocol_sort_rows(run_key, run)
                place = 0
                for r in rows:
                    st_u = (r.get("status") or "").strip().upper()
                    if r.get("finish_f") is not None and st_u not in ("DNS",):
                        place += 1
                        r["place"] = str(place)
                    else:
                        r["place"] = ""

                doc.add_heading("Протокол", level=1)
                if ev:
                    doc.add_paragraph(f"Дисциплина: {ev}")

                date_line = date_line_global
                if not date_line:
                    st = str(run.get("start_time") or run.get("start") or "").strip()
                    if st:
                        try:
                            s2 = st.replace("Z", "+00:00")
                            dt = datetime.datetime.fromisoformat(s2)
                            date_line = fmt_ru_long_date(dt.date())
                        except Exception:
                            pass
                if date_line:
                    doc.add_paragraph(date_line)
                if conditions_global:
                    doc.add_paragraph(conditions_global)
                if dm is not None:
                    doc.add_paragraph(f"Дистанция: {dm} м")
                if rnd:
                    doc.add_paragraph(f"Раунд: {rnd}")
                if cat:
                    doc.add_paragraph(f"Категория: {cat}")
                doc.add_paragraph(f"Заезд: {run_key}")
                if start:
                    doc.add_paragraph(f"Старт: {start}")
                if judge:
                    doc.add_paragraph(f"Главный судья: {judge}")
                if sec:
                    doc.add_paragraph(f"Секретарь: {sec}")
                if notes:
                    doc.add_paragraph(f"Примечание: {notes}")

                cols = ["Место", "№", "Имя", "Стр/Гор", "Финиш", "Статус"]
                if include_splits and split_ids:
                    cols.append("Отсечки")

                table = doc.add_table(rows=1, cols=len(cols))
                hdr = table.rows[0].cells
                for j, c in enumerate(cols):
                    hdr[j].text = c

                for r in rows:
                    cells = table.add_row().cells
                    cells[0].text = r.get("place", "")
                    cells[1].text = r.get("bib", "")
                    cells[2].text = r.get("name", "")
                    cells[3].text = r.get("country", "")
                    cells[4].text = r.get("finish", "")
                    cells[5].text = r.get("status", "")
                    if include_splits and split_ids:
                        sp = []
                        splits = r.get("splits") if isinstance(r.get("splits"), dict) else {}
                        for sid in split_ids:
                            if str(sid) in splits:
                                sp.append(f"S{sid}:{fmt_time(splits.get(str(sid)))}")
                        cells[-1].text = " ".join(sp)

            base = f"{proto_type or 'protocol'}_{run_key}" if proto_type else f"protocol_{run_key}"
            filename = _safe_name(base) + ".docx"
            out_path = str(Path(folder) / filename)
            try:
                doc.save(out_path)
                saved += 1
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить {out_path}:\n{e}")
                return

        messagebox.showinfo("Готово", f"Сохранено файлов: {saved}\nПапка: {folder}")

    def connect(self):
        if self._thr and self._thr.is_alive():
            return
        host = self.host_var.get().strip() or "127.0.0.1"
        try:
            port = int(self.port_var.get().strip())
        except Exception:
            messagebox.showerror("Ошибка", "Неверный port")
            return

        self._stop_evt.clear()
        self._thr = TcpClientThread(host, port, self._q, self._stop_evt)
        self._thr.start()

    def disconnect(self):
        self._stop_evt.set()

    def _set_status(self, text: str, ok: bool):
        self.status_var.set(text)
        color = self.colors["ok"] if ok else self.colors["danger"]
        try:
            self.status_dot.itemconfigure(self._dot_id, fill=color, outline=color)
        except Exception:
            pass
        self._connected = ok

    def _pump(self):
        try:
            while True:
                item = self._q.get_nowait()
                k = item.get("kind")
                if k == "status":
                    self._set_status(str(item.get("text", "")), bool(item.get("ok")))
                elif k == "error":
                    self._append_log("ERROR: " + str(item.get("text", "")))
                elif k == "raw":
                    self._append_log(str(item.get("text", "")))
                elif k == "msg":
                    raw = item.get("raw")
                    if raw:
                        self._append_log(raw)
                    self._handle_message(item.get("data") or {})
        except queue.Empty:
            pass

        if (not self._connected) and self.auto_reconnect_var.get():
            if (time.time() - self._last_state_ts) > 2.0:
                if not (self._thr and self._thr.is_alive()):
                    self._last_state_ts = time.time()
                    self.connect()

        self.after(80, self._pump)

    def _handle_message(self, msg: Dict[str, Any]):
        t = msg.get("type")
        if t == "hello":
            return
        if t == "state":
            st = msg.get("state")
            if isinstance(st, dict):
                self.state = st
                self._last_state_ts = time.time()
                self._refresh_views()
                return
    def _update_info_bar(self):
        try:
            runs = self.state.get("runs", {}) if isinstance(self.state.get("runs"), dict) else {}
            n_runs = len(runs)
            n_ath = 0
            n_fin = 0
            for _rk, run in runs.items():
                if not isinstance(run, dict):
                    continue
                ath = run.get("athletes")
                if isinstance(ath, dict):
                    n_ath += len(ath)
                    for _b, a in ath.items():
                        if isinstance(a, dict) and a.get("finish") is not None:
                            n_fin += 1
            t = time.strftime("%H:%M:%S")
            self.info_var.set(f"{n_runs} заезд(ов), {n_fin}/{n_ath} финиш, {t}")
        except Exception:
            try:
                self.info_var.set("")
            except Exception:
                pass


    def _refresh_views(self):
        self._render_runs()
        self._render_runs_table_text()
        self._update_export_preview()
        self._update_protocol_preview()


    def _treeview_flat_order(self, tv: ttk.Treeview) -> List[str]:
        out: List[str] = []

        def walk(parent: str):
            for iid in tv.get_children(parent):
                out.append(str(iid))
                walk(iid)

        walk("")
        return out
    def _copy_treeview_rows(self, tv: ttk.Treeview, iids: List[str]):
        iids = [str(x) for x in (iids or []) if str(x)]
        if not iids:
            return

        col_count = 1 + len(tv["columns"])
        lines: List[str] = []
        for iid in iids:
            item = tv.item(iid)
            text = str(item.get("text") or "")
            values = item.get("values") or ()
            row = [text] + [str(v) for v in values]
            if len(row) < col_count:
                row += [""] * (col_count - len(row))
            lines.append("\t".join(row))

        out = "\n".join(lines).rstrip() + "\n"
        try:
            self.clipboard_clear()
            self.clipboard_append(out)
        except Exception:
            pass

    def _copy_treeview_selection(self, tv: ttk.Treeview):
        sel = list(tv.selection())
        if not sel:
            return

        order = self._treeview_flat_order(tv)
        pos = {iid: i for i, iid in enumerate(order)}
        sel = [str(iid) for iid in sel]
        sel.sort(key=lambda iid: pos.get(str(iid), 10**9))
        self._copy_treeview_rows(tv, sel)

    def _treeview_cell_value(self, tv: ttk.Treeview, iid: str, col: str) -> str:
        try:
            item = tv.item(iid)
        except Exception:
            return ""
        if col == "#0":
            return str(item.get("text") or "")
        try:
            idx = int(col.replace("#", "")) - 1
        except Exception:
            return ""
        vals = item.get("values") or ()
        if idx < 0 or idx >= len(vals):
            return ""
        v = vals[idx]
        return "" if v is None else str(v)

    def _runs_tree_record_cell(self, evt):
        try:
            iid = self.runs_tv.identify_row(evt.y)
            col = self.runs_tv.identify_column(evt.x)
            if iid:
                self._runs_active_iid = str(iid)
                self._runs_active_col = str(col)
        except Exception:
            pass

    def _copy_runs_tree_cell(self):
        iid = self._runs_active_iid
        col = self._runs_active_col
        if not iid or not col:
            return
        txt = self._treeview_cell_value(self.runs_tv, iid, col)
        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
        except Exception:
            pass
    def _runs_is_run_row(self, iid: str) -> bool:
        iid = str(iid)
        if not iid or ":" in iid:
            return False
        try:
            return len(self.runs_tv.get_children(iid)) > 0
        except Exception:
            return False

    def _runs_selected_run_keys(self) -> List[str]:
        sel = list(self.runs_tv.selection())
        run_keys: List[str] = []
        seen = set()
        for iid in sel:
            iid = str(iid)
            rk = iid
            if ":" in iid:
                try:
                    parent = self.runs_tv.parent(iid)
                    rk = parent or iid.split(":", 1)[0]
                except Exception:
                    rk = iid.split(":", 1)[0]
            if rk and rk not in seen and self._runs_is_run_row(rk):
                seen.add(rk)
                run_keys.append(rk)

        # keep in screen order
        ordered = []
        try:
            root_children = [str(x) for x in self.runs_tv.get_children("")]
            for rk in root_children:
                if rk in seen:
                    ordered.append(rk)
        except Exception:
            ordered = run_keys
        return ordered

    def _runs_expand_iids(self, sel_iids: List[str], include_run_rows: bool = True) -> List[str]:
        tv = self.runs_tv
        want = set()

        for iid in sel_iids:
            iid = str(iid)
            if not iid:
                continue

            if include_run_rows:
                want.add(iid)

            # expand run rows to their descendants
            if self._runs_is_run_row(iid):
                want.add(iid)
                stack = [str(x) for x in tv.get_children(iid)]
                while stack:
                    cur = stack.pop(0)
                    want.add(cur)
                    for ch in tv.get_children(cur):
                        stack.append(str(ch))

        # order by visual traversal
        order = self._treeview_flat_order(tv)
        pos = {iid: i for i, iid in enumerate(order)}
        out = [iid for iid in sorted(want, key=lambda x: pos.get(x, 10**9)) if iid]
        return out

    def _runs_select_blocks(self):
        run_keys = self._runs_selected_run_keys()
        if not run_keys:
            return
        iids = self._runs_expand_iids(run_keys, include_run_rows=True)
        try:
            self.runs_tv.selection_set(iids)
        except Exception:
            pass

    def _copy_runs_selected_runs_full(self):
        run_keys = self._runs_selected_run_keys()
        if not run_keys:
            return
        iids = self._runs_expand_iids(run_keys, include_run_rows=True)
        self._copy_treeview_rows(self.runs_tv, iids)

    def _copy_runs_tree_selection(self, _evt=None):
        sel = list(self.runs_tv.selection())
        if self._runs_active_iid and self._runs_active_col:
            if not sel or (len(sel) == 1 and str(sel[0]) == str(self._runs_active_iid)):
                self._copy_runs_tree_cell()
                return "break"

        if not sel:
            return "break"

        # if at least one run row is selected -> copy whole run(s) with athletes
        if any(self._runs_is_run_row(iid) for iid in sel):
            run_keys = self._runs_selected_run_keys()
            if not run_keys:
                # fallback: expand raw selection
                iids = self._runs_expand_iids([str(x) for x in sel], include_run_rows=True)
                self._copy_treeview_rows(self.runs_tv, iids)
            else:
                self._copy_runs_selected_runs_full()
            return "break"

        self._copy_treeview_selection(self.runs_tv)
        return "break"

    def _runs_tree_context_menu(self, evt):
        try:
            iid = self.runs_tv.identify_row(evt.y)
            col = self.runs_tv.identify_column(evt.x)
            if iid:
                self._runs_active_iid = str(iid)
                self._runs_active_col = str(col)
            if iid and iid not in self.runs_tv.selection():
                self.runs_tv.selection_add(iid)
        except Exception:
            pass

        m = tk.Menu(self, tearoff=0)
        m.add_command(label="Копировать ячейку", command=self._copy_runs_tree_cell)
        m.add_command(label="Копировать строки (как выделено)", command=lambda: self._copy_treeview_selection(self.runs_tv))
        m.add_separator()
        m.add_command(label="Выделить заезд(ы) целиком", command=self._runs_select_blocks)
        m.add_command(label="Копировать заезд(ы) целиком", command=self._copy_runs_selected_runs_full)
        try:
            m.tk_popup(evt.x_root, evt.y_root)
        finally:
            try:
                m.grab_release()
            except Exception:
                pass




    def _ensure_runs_columns(self, split_ids: List[str]):
        if not hasattr(self, "runs_tv") or not self.runs_tv:
            return

        split_ids = [str(sid) for sid in split_ids if str(sid).strip()]
        desired = ["cat", "dist", "start", "place", "country"] + [f"S{sid}" for sid in split_ids] + ["finish", "status"]
        if desired == self._runs_cols and split_ids == self._runs_split_ids:
            return

        self._runs_cols = list(desired)
        self._runs_split_ids = list(split_ids)

        self.runs_tv["columns"] = tuple(self._runs_cols)

        self.runs_tv.heading("#0", text="Имя")
        self.runs_tv.heading("cat", text="Кат.")
        self.runs_tv.heading("dist", text="Дист., м")
        self.runs_tv.heading("start", text="Старт")
        self.runs_tv.heading("place", text="Место")
        self.runs_tv.heading("country", text="Стр/Гор")

        for sid in split_ids:
            cid = f"S{sid}"
            self.runs_tv.heading(cid, text=cid)

        self.runs_tv.heading("finish", text="Финиш")
        self.runs_tv.heading("status", text="Статус")

        self.runs_tv.column("#0", width=360, anchor="w")
        self.runs_tv.column("cat", width=120, anchor="center")
        self.runs_tv.column("dist", width=90, anchor="center")
        self.runs_tv.column("start", width=140, anchor="center")
        self.runs_tv.column("place", width=70, anchor="center")
        self.runs_tv.column("country", width=140, anchor="center")

        for sid in split_ids:
            cid = f"S{sid}"
            self.runs_tv.column(cid, width=110, anchor="center")

        self.runs_tv.column("finish", width=110, anchor="center")
        self.runs_tv.column("status", width=110, anchor="center")


    def _splits_to_map(self, splits: Any) -> Dict[str, str]:
        out: Dict[str, str] = {}
        if isinstance(splits, dict):
            for k, v in splits.items():
                ks = str(k).strip()
                if ks:
                    out[ks] = fmt_time(v)
            return out
        if isinstance(splits, (list, tuple)):
            for i, v in enumerate(splits):
                out[str(i + 1)] = fmt_time(v)
            return out
        return out


    def _render_runs(self):
        if not hasattr(self, "runs_tv"):
            return

        for iid in self.runs_tv.get_children(""):
            self.runs_tv.delete(iid)

        items = self._runs_in_display_order()
        seen_split_ids: set = set()
        for _rk, _run in items:
            for _sid in self._run_split_ids(_run):
                seen_split_ids.add(str(_sid))
        global_split_ids = sorted(seen_split_ids, key=split_sort_key)
        self._ensure_runs_columns(global_split_ids)
        flt = self.run_filter_var.get().strip().lower()

        idx = 0
        for run_key, run in items:
            start = str(run.get("start_time") or "").strip()
            ev = str(run.get("event") or run.get("discipline") or run.get("name") or "").strip()
            rnd = str(run.get("round") or run.get("phase") or run.get("heat") or "").strip()
            cat = self.settings.run_categories.get(run_key, "") or ""

            dm = run_distance_m(run)
            dm_s = str(dm) if dm is not None else ""

            rows = self._protocol_sort_rows(run_key, run)

            place = 0
            for r in rows:
                st_u = (r.get("status") or "").strip().upper()
                if r.get("finish_f") is not None and st_u not in ("DNS",):
                    place += 1
                    r["place"] = str(place)
                else:
                    r["place"] = ""

            split_ids = self._runs_split_ids

            def _ath_hay(r: Dict[str, Any]) -> str:
                return f"{r.get('bib','')} {r.get('name','')} {r.get('country','')} {r.get('finish','')} {r.get('status','')}".lower()

            run_hay = f"{run_key} {ev} {rnd} {cat} {start} {dm_s}".lower()

            if flt:
                if flt in run_hay:
                    shown = rows
                else:
                    shown = [r for r in rows if flt in _ath_hay(r)]
                if not shown:
                    continue
            else:
                shown = rows

            ath_n = len(rows)
            fin_n = 0
            for r in rows:
                if r.get("finish_f") is not None:
                    fin_n += 1

            run_text = f"Заезд {run_key}"
            if ev:
                run_text += f" — {ev}"
            if rnd:
                run_text += f" ({rnd})"
            if ath_n:
                run_text += f"  [{fin_n}/{ath_n}]"

            run_has_warn = False
            run_has_missing = False
            tag = "run_even" if (idx % 2 == 0) else "run_odd"
            self.runs_tv.insert("", "end",
                                iid=str(run_key),
                                text=run_text,
                                values=tuple([cat, dm_s, start, "", ""] + [""] * len(split_ids) + ["", ""]),
                                tags=(tag,))

            for r in shown:
                bib = str(r.get("bib") or "").strip()
                name = str(r.get("name") or "").strip()
                country = str(r.get("country") or "").strip()
                finish = str(r.get("finish") or "").strip()
                status = str(r.get("status") or "").strip()

                spm = self._splits_to_map(r.get("splits"))
                sp_vals = [spm.get(str(sid), "") for sid in split_ids]

                if name:
                    child_text = f"№{bib} {name}".strip()
                else:
                    child_text = f"№{bib}".strip()

                warn = False
                try:
                    if (not finish) and (not status):
                        warn = True
                    if str(status).strip().startswith("#"):
                        warn = True
                    if finish in ("0.000", "0,000", "0:00.000", "0:00,000") and (not status):
                        warn = True
                except Exception:
                    pass
                missing = False
                try:
                    if cat and (self.roster.get_entry(cat, bib) is None):
                        missing = True
                except Exception:
                    pass
                child_tags = ["ath"]
                if missing:
                    child_tags.append("miss")
                if warn:
                    child_tags.append("warn")
                run_has_warn = run_has_warn or warn
                run_has_missing = run_has_missing or missing

                cid = f"{run_key}:{bib}"
                self.runs_tv.insert(str(run_key), "end",
                                    iid=cid,
                                    text=child_text,
                                    values=tuple(["", "", "", str(r.get("place", "") or ""), country] + sp_vals + [finish, status]),
                                    tags=tuple(child_tags))

            if run_has_warn or run_has_missing:
                try:
                    cur_tags = list(self.runs_tv.item(str(run_key), "tags") or ())
                    if "run_warn" not in cur_tags:
                        cur_tags.append("run_warn")
                    self.runs_tv.item(str(run_key), tags=tuple(cur_tags))
                except Exception:
                    pass
            try:
                self.runs_tv.item(str(run_key), open=True)
            except Exception:
                pass

            idx += 1

        cur = self.selected_run_key or self.state.get("current_key")
        if cur and str(cur) in self.runs_tv.get_children(""):
            try:
                self.runs_tv.selection_set(str(cur))
                self.runs_tv.see(str(cur))
            except Exception:
                pass


    def _on_run_select(self, _evt=None):
        sel = self.runs_tv.selection()
        if not sel:
            self.selected_run_key = None
            self.run_category_var.set("")
            self._update_export_preview()
            self._update_protocol_preview()
            return

        iid = str(sel[0])
        run_key = iid
        if ":" in iid:
            try:
                parent = self.runs_tv.parent(iid)
                if parent:
                    run_key = parent
                else:
                    run_key = iid.split(":", 1)[0]
            except Exception:
                run_key = iid.split(":", 1)[0]

        self.selected_run_key = run_key

        cat = self.settings.run_categories.get(run_key, "")
        if cat:
            self.run_category_var.set(cat)
        else:
            if self.category_var.get().strip():
                self.run_category_var.set(self.category_var.get().strip())

        self._update_export_preview()
        self._update_protocol_preview()


    def _on_run_category_changed(self):
        run_key = self.selected_run_key
        if not run_key:
            return
        cat = self.run_category_var.get().strip()
        if not cat:
            return
        self.settings.run_categories[run_key] = cat
        self.category_var.set(cat)
        self._render_runs()
        self._render_roster()
        self._update_export_preview()
        self._update_protocol_preview()

    def _run_split_ids(self, run: Dict[str, Any]) -> List[str]:
        ath = run.get("athletes")
        ids = set()
        if isinstance(ath, dict):
            for _bib, a in ath.items():
                if not isinstance(a, dict):
                    continue
                sp = a.get("splits")
                if isinstance(sp, dict):
                    for k in sp.keys():
                        ks = str(k).strip()
                        if ks:
                            ids.add(ks)
                elif isinstance(sp, (list, tuple)):
                    for i in range(len(sp)):
                        ids.add(str(i + 1))
        return sorted(ids, key=split_sort_key)

    def _rebuild_ath_tree(self, split_ids: List[str]):
        for child in self.ath_container.winfo_children():
            child.destroy()

        cols = ["bib", "name", "country"]
        if self.show_distance_var.get():
            cols.append("dist")
        cols += [f"S{sid}" for sid in split_ids] + ["finish", "status"]

        self._ath_cols = cols
        self.ath_tv = ttk.Treeview(self.ath_container, columns=cols, show="headings")
        vsb = ttk.Scrollbar(self.ath_container, orient="vertical", command=self.ath_tv.yview)
        hsb = ttk.Scrollbar(self.ath_container, orient="horizontal", command=self.ath_tv.xview)
        self.ath_tv.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.ath_tv.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        self.ath_container.rowconfigure(0, weight=1)
        self.ath_container.columnconfigure(0, weight=1)

        self.ath_tv.heading("bib", text="№")
        self.ath_tv.column("bib", width=90, anchor="center")
        self.ath_tv.heading("name", text="Имя")
        self.ath_tv.column("name", width=360, anchor="w")
        self.ath_tv.heading("country", text="Стр/Гор")
        self.ath_tv.column("country", width=120, anchor="center")

        if self.show_distance_var.get():
            self.ath_tv.heading("dist", text="Дист.")
            self.ath_tv.column("dist", width=90, anchor="center")

        for sid in split_ids:
            cid = f"S{sid}"
            self.ath_tv.heading(cid, text=f"S{sid}")
            self.ath_tv.column(cid, width=125, anchor="center")

        self.ath_tv.heading("finish", text="Финиш")
        self.ath_tv.column("finish", width=140, anchor="center")
        self.ath_tv.heading("status", text="Статус")
        self.ath_tv.column("status", width=160, anchor="center")

        self.ath_tv.tag_configure("odd", background=self.colors["odd"])
        self.ath_tv.tag_configure("even", background=self.colors["even"])

    def _effective_meta(self, run_key: str, athlete: Dict[str, Any]) -> Tuple[str, str]:
        bib = safe_int_str(athlete.get("bib") or "")
        if not bib:
            return "", ""
        srv_name = str(athlete.get("name") or "").strip()
        srv_country = str(athlete.get("country") or "").strip().upper()

        cat = self.settings.run_categories.get(run_key, "") or self.run_category_var.get().strip() or self.category_var.get().strip()
        if not cat:
            return srv_name, srv_country

        entry = self.roster.get_entry(cat, bib)
        if not entry:
            return srv_name, srv_country

        if self.override_names_var.get() or not srv_name:
            name = entry.get("name", "") or srv_name
        else:
            name = srv_name

        if self.override_names_var.get() or not srv_country:
            country = entry.get("country", "") or srv_country
        else:
            country = srv_country

        return name, country

    def _effective_meta_full(self, run_key: str, athlete: Dict[str, Any]) -> Dict[str, str]:
        """name/country + optional fields for protocols (org, dob, rank, region)."""
        bib = safe_int_str(athlete.get("bib") or "")
        srv: Dict[str, str] = {
            "bib": bib,
            "name": str(athlete.get("name") or "").strip(),
            "country": str(athlete.get("country") or "").strip().upper(),
            "org": str(athlete.get("org") or athlete.get("club") or athlete.get("team") or athlete.get("organization") or "").strip(),
            "dob": str(athlete.get("dob") or athlete.get("birthdate") or athlete.get("birth_date") or athlete.get("date_of_birth") or "").strip(),
            "rank": str(athlete.get("rank") or athlete.get("category") or athlete.get("class") or "").strip(),
            "region": str(athlete.get("region") or athlete.get("city") or athlete.get("place") or "").strip(),
        }
        if not bib:
            return srv

        cat = self.settings.run_categories.get(run_key, "") or self.run_category_var.get().strip() or self.category_var.get().strip()
        if not cat:
            return srv

        entry = self.roster.get_entry(cat, bib)
        if not entry:
            return srv

        use_override = bool(self.override_names_var.get())

        def pick(k: str, default: str) -> str:
            v = str(entry.get(k) or "").strip()
            if use_override or not default:
                return v or default
            return default

        srv["name"] = pick("name", srv.get("name", ""))
        srv["country"] = pick("country", srv.get("country", "")).upper()

        for k in ("org", "dob", "rank", "region"):
            srv[k] = pick(k, srv.get(k, ""))

        return srv

    def _render_athletes(self, run_key: Optional[str]):
        if not self.ath_tv:
            return
        for iid in self.ath_tv.get_children():
            self.ath_tv.delete(iid)

        if not run_key:
            self.run_info_var.set("—")
            return

        runs = self.state.get("runs", {}) if isinstance(self.state.get("runs"), dict) else {}
        run = runs.get(run_key) if isinstance(runs.get(run_key), dict) else None
        if not run:
            self.run_info_var.set("—")
            return

        split_ids = self._run_split_ids(run)
        desired_cols = ["bib", "name", "country"]
        if self.show_distance_var.get():
            desired_cols.append("dist")
        desired_cols += [f"S{sid}" for sid in split_ids] + ["finish", "status"]
        if desired_cols != self._ath_cols:
            self._rebuild_ath_tree(split_ids=split_ids)
            self.ath_tv.bind("<Double-1>", self._on_athlete_double_click)

        start = run.get("start_time") or "—"
        ath = run.get("athletes") or {}
        ath_n = len(ath) if isinstance(ath, dict) else 0
        fin_n = 0
        if isinstance(ath, dict):
            for _bib, a in ath.items():
                if isinstance(a, dict) and a.get("finish") is not None:
                    fin_n += 1

        cat = self.settings.run_categories.get(run_key, "") or "—"
        self.run_info_var.set(f"{run_key}   кат: {cat}   старт: {start}   участников: {ath_n}   финиш: {fin_n}")

        flt = self.ath_filter_var.get().strip().lower()

        order = run.get("bib_order")
        if isinstance(order, list) and order:
            bibs = [safe_int_str(b) for b in order if safe_int_str(b)]
        else:
            bibs = [safe_int_str(b) for b in (ath.keys() if isinstance(ath, dict) else [])]
        bibs = [b for b in bibs if b in ath]

        idx = 0
        for bib in bibs:
            a = ath.get(bib)
            if not isinstance(a, dict):
                continue

            eff_name, eff_country = self._effective_meta(run_key, a)

            splits = a.get("splits") or {}
            finish = a.get("finish")
            status = a.get("status") or ""

            dist = ""
            if self.show_distance_var.get():
                dm = athlete_distance_m(a)
                if dm is None and a.get("finish") is not None:
                    dm = run_distance_m(run)
                if dm is not None:
                    dist = fmt_dist(dm)
                else:
                    split_count = len(splits) if isinstance(splits, dict) else 0
                    dist = f"{split_count * DIST_PER_SPLIT_M}м" if split_count > 0 else  ""

            if flt:
                hay = f"{bib} {eff_name} {eff_country} {status}".lower()
                if flt not in hay:
                    continue

            row = [bib, eff_name, eff_country]
            if self.show_distance_var.get():
                row.append(dist)

            for sid in split_ids:
                v = ""
                if isinstance(splits, dict):
                    v = fmt_time(splits.get(str(sid)))
                row.append(v)

            row.append(fmt_time(finish))
            row.append(str(status))

            tag = "even" if (idx % 2 == 0) else "odd"
            self.ath_tv.insert("", "end", iid=f"{run_key}:{bib}", values=row, tags=(tag,))
            idx += 1

    def _render_roster(self):
        for iid in self.roster_tv.get_children():
            self.roster_tv.delete(iid)

        cat = self.category_var.get().strip()
        if not cat:
            return

        mp = self.roster.data.get(cat, {})
        bibs = sorted(mp.keys(), key=split_sort_key)

        idx = 0
        for bib in bibs:
            meta = mp.get(bib) or {}
            name = meta.get("name", "")
            country = meta.get("country", "")
            tag = "even" if (idx % 2 == 0) else "odd"
            self.roster_tv.insert("", "end", iid=f"{cat}:{bib}", values=(bib, name, country), tags=(tag,))
            idx += 1

    def _save_roster(self):
        self.roster.save()
        self._append_log("ROSTER: saved")

    def _clear_roster_category(self):
        cat = self.category_var.get().strip()
        if not cat:
            return
        if messagebox.askyesno("Очистить", f"Очистить состав категории: {cat}?"):
            self.roster.clear_category(cat)
            self.roster.save()
            self._render_roster()
            self._refresh_views()

    def _edit_roster_entry(self, cat: str, bib: str, name: str, country: str) -> bool:
        d = tk.Toplevel(self)
        d.title("Запись состава")
        d.resizable(False, False)

        v_bib = tk.StringVar(value=bib)
        v_name = tk.StringVar(value=name)
        v_country = tk.StringVar(value=country)

        frm = ttk.Frame(d)
        frm.pack(padx=12, pady=12, fill="x")

        ttk.Label(frm, text=f"Категория: {cat}", style="Muted.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))

        ttk.Label(frm, text="№").grid(row=1, column=0, sticky="w")
        ttk.Entry(frm, textvariable=v_bib, width=10).grid(row=1, column=1, sticky="w")

        ttk.Label(frm, text="Имя").grid(row=2, column=0, sticky="w", pady=(8, 0))
        e_n = ttk.Entry(frm, textvariable=v_name, width=42)
        e_n.grid(row=2, column=1, sticky="w", pady=(8, 0))

        ttk.Label(frm, text="Страна/Город").grid(row=3, column=0, sticky="w", pady=(8, 0))
        ttk.Entry(frm, textvariable=v_country, width=14).grid(row=3, column=1, sticky="w", pady=(8, 0))

        out = {"ok": False}

        def ok():
            bb = safe_int_str(v_bib.get())
            if not bb:
                messagebox.showerror("Ошибка", "Неверный номер")
                return
            self.roster.set_entry(cat, bb, v_name.get(), v_country.get())
            self.roster.save()
            out["ok"] = True
            d.destroy()

        def delete():
            bb = safe_int_str(v_bib.get())
            if not bb:
                return
            if messagebox.askyesno("Удалить", f"Удалить запись №{bb} из {cat}?"):
                self.roster.delete_entry(cat, bb)
                self.roster.save()
                out["ok"] = True
                d.destroy()

        def cancel():
            d.destroy()

        btns = ttk.Frame(d)
        btns.pack(padx=12, pady=(0, 12), fill="x")
        ttk.Button(btns, text="Отмена", command=cancel).pack(side="right")
        ttk.Button(btns, text="OK", style="Accent.TButton", command=ok).pack(side="right", padx=(0, 8))
        ttk.Button(btns, text="Удалить", command=delete).pack(side="left")

        d.grab_set()
        e_n.focus_set()
        d.wait_window()
        return out["ok"]

    def _on_roster_double_click(self, _evt=None):
        sel = self.roster_tv.selection()
        if not sel:
            return
        iid = sel[0]
        try:
            cat, bib = iid.split(":", 1)
        except Exception:
            return
        mp = self.roster.data.get(cat, {})
        meta = mp.get(bib, {})
        if self._edit_roster_entry(cat, bib, meta.get("name", ""), meta.get("country", "")):
            self._render_roster()
            self._refresh_views()


    def _on_athlete_double_click(self, _evt=None):
        tv = self.ath_tv if getattr(self, "ath_tv", None) else getattr(self, "runs_tv", None)
        if not tv:
            return

        iid = None
        if _evt is not None and hasattr(tv, "identify_row"):
            try:
                iid = tv.identify_row(getattr(_evt, "y", 0))
            except Exception:
                iid = None

        if not iid:
            sel = tv.selection()
            if not sel:
                return
            iid = sel[0]

        iid = str(iid)
        if ":" not in iid:
            return

        try:
            run_key, bib = iid.split(":", 1)
        except Exception:
            return

        cat = self.settings.run_categories.get(run_key, "") or self.run_category_var.get().strip() or self.category_var.get().strip()
        if not cat:
            messagebox.showerror("Ошибка", "Сначала назначь категорию заезда")
            return

        name, country = "", ""
        entry = self.roster.get_entry(cat, bib)
        if entry:
            name, country = entry.get("name", ""), entry.get("country", "")
        else:
            runs = self.state.get("runs", {}) if isinstance(self.state.get("runs"), dict) else {}
            run = runs.get(run_key, {}) if isinstance(runs.get(run_key), dict) else {}
            ath = run.get("athletes") or {}
            a = ath.get(bib) if isinstance(ath, dict) else None
            if isinstance(a, dict):
                name = str(a.get("name") or "").strip()
                country = str(a.get("country") or "").strip().upper()

        if not name and not country:
            name = ""
            country = ""

        d = AthleteDialog(self, bib=bib, name=name, country=country)
        d.wait_window()
        if not d.result:
            return

        self.roster.set_entry(cat, bib, d.result.get("name", ""), d.result.get("country", ""))
        self._save_roster()
        self._render_runs()
        self._render_roster()
        self._update_export_preview()
        self._update_protocol_preview()

    def _load_roster_for_current_category(self):
        cat = self.run_category_var.get().strip() or self.category_var.get().strip()
        if not cat:
            messagebox.showerror("Ошибка", "Сначала выбери категорию")
            return
        self.category_var.set(cat)
        self._load_roster_dialog(forced_category=cat)

    def _load_roster_dialog(self, forced_category: Optional[str] = None):
        ft = [("Excel", "*.xlsx")] if load_workbook else []
        ft += [("CSV", "*.csv"), ("All", "*.*")]
        path = filedialog.askopenfilename(title="Загрузить состав", filetypes=ft)
        if not path:
            return
        try:
            loaded, cats = self._load_roster_file(path, forced_category=forced_category)
            self.roster.save()
            self._append_log(f"ROSTER: loaded {loaded} records")
            if cats:
                self._append_log("ROSTER categories: " + ", ".join(cats))
            self._render_roster()
            self._refresh_views()
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def _load_roster_file(self, path: str, forced_category: Optional[str]) -> Tuple[int, List[str]]:
        p = Path(path)
        ext = p.suffix.lower()
        if ext == ".xlsx":
            if load_workbook is None:
                raise RuntimeError("openpyxl не установлен")
            return self._load_roster_xlsx(p, forced_category)
        return self._load_roster_csv(p, forced_category)

    def _load_roster_xlsx(self, p: Path, forced_category: Optional[str]) -> Tuple[int, List[str]]:
        wb = load_workbook(str(p), data_only=True)
        ws = wb.active

        cat_default = forced_category or self.category_var.get().strip()
        if not cat_default:
            cat_default = self.settings.categories[0] if self.settings.categories else "Категория"

        header = [str(x).strip().lower() if x is not None else "" for x in next(ws.iter_rows(min_row=1, max_row=1, values_only=True))]
        has_header = any(h in ("bib", "номер", "№", "name", "имя", "country", "страна", "категория", "category") for h in header)
        start_row = 2 if has_header else 1

        def idx_of(names: List[str]) -> Optional[int]:
            for i, h in enumerate(header):
                if h in names:
                    return i
            return None

        i_bib = idx_of(["bib", "номер", "№"]) or 0
        i_name = idx_of(["name", "имя", "фио"]) or 1
        i_country = idx_of(["country", "страна", "город"]) or 2
        i_cat = idx_of(["category", "категория"])

        i_org = idx_of(["org", "organization", "организация", "школа", "клуб", "команда"])
        i_dob = idx_of(["dob", "birthdate", "date_of_birth", "дата рождения", "д.р.", "др"])
        i_rank = idx_of(["rank", "разряд", "спорт разряд", "спортразряд"])
        i_region = idx_of(["region", "регион", "область", "город/регион"])

        touched = set()
        n = 0
        for r in ws.iter_rows(min_row=start_row, values_only=True):
            if not r:
                continue
            bib = safe_int_str(r[i_bib] if i_bib < len(r) else "")
            if not bib:
                continue
            name = str(r[i_name]).strip() if i_name < len(r) and r[i_name] is not None else ""
            country = str(r[i_country]).strip().upper() if i_country < len(r) and r[i_country] is not None else ""
            cat = None
            if i_cat is not None and i_cat < len(r) and r[i_cat] is not None:
                cat = str(r[i_cat]).strip()
            if not cat:
                cat = cat_default
            org = str(r[i_org]).strip() if i_org is not None and i_org < len(r) and r[i_org] is not None else ""
            dob = str(r[i_dob]).strip() if i_dob is not None and i_dob < len(r) and r[i_dob] is not None else ""
            rank = str(r[i_rank]).strip() if i_rank is not None and i_rank < len(r) and r[i_rank] is not None else ""
            region = str(r[i_region]).strip() if i_region is not None and i_region < len(r) and r[i_region] is not None else ""
            self.roster.set_entry(cat, bib, name, country, org=org, dob=dob, rank=rank, region=region)
            touched.add(cat)
            n += 1

        return n, sorted(touched)

    def _load_roster_csv(self, p: Path, forced_category: Optional[str]) -> Tuple[int, List[str]]:
        cat_default = forced_category or self.category_var.get().strip()
        if not cat_default:
            cat_default = self.settings.categories[0] if self.settings.categories else "Категория"

        touched = set()
        n = 0

        text = p.read_text(encoding="utf-8", errors="ignore")
        first_line = text.splitlines()[0] if text.splitlines() else ""
        delim = ";" if first_line.count(";") >= first_line.count(",") else ","

        rows = list(csv.reader(text.splitlines(), delimiter=delim))
        if not rows:
            return 0, []

        header = [str(x).strip().lower() for x in rows[0]]
        has_header = any(h in ("bib", "номер", "№", "name", "имя", "фио", "country", "страна", "город", "category", "категория") for h in header)

        def idx_of(names: List[str]) -> Optional[int]:
            for i, h in enumerate(header):
                if h in names:
                    return i
            return None

        i_bib = idx_of(["bib", "номер", "№"]) if has_header else 0
        i_name = idx_of(["name", "имя", "фио"]) if has_header else 1
        i_country = idx_of(["country", "страна", "город"]) if has_header else 2
        i_cat = idx_of(["category", "категория"]) if has_header else None

        i_org = idx_of(["org", "organization", "организация", "школа", "клуб", "команда"]) if has_header else None
        i_dob = idx_of(["dob", "birthdate", "date_of_birth", "дата рождения", "д.р.", "др"]) if has_header else None
        i_rank = idx_of(["rank", "разряд", "спорт разряд", "спортразряд"]) if has_header else None
        i_region = idx_of(["region", "регион", "область", "город/регион"]) if has_header else None

        start = 1 if has_header else 0

        for r in rows[start:]:
            if not r:
                continue
            bib = safe_int_str(r[i_bib] if i_bib is not None and i_bib < len(r) else "")
            if not bib:
                continue
            name = str(r[i_name]).strip() if i_name is not None and i_name < len(r) else ""
            country = str(r[i_country]).strip().upper() if i_country is not None and i_country < len(r) else ""
            cat = ""
            if i_cat is not None and i_cat < len(r):
                cat = str(r[i_cat]).strip()
            if not cat:
                cat = cat_default
            org = str(r[i_org]).strip() if i_org is not None and i_org < len(r) else ""
            dob = str(r[i_dob]).strip() if i_dob is not None and i_dob < len(r) else ""
            rank = str(r[i_rank]).strip() if i_rank is not None and i_rank < len(r) else ""
            region = str(r[i_region]).strip() if i_region is not None and i_region < len(r) else ""
            self.roster.set_entry(cat, bib, name, country, org=org, dob=dob, rank=rank, region=region)
            touched.add(cat)
            n += 1

        return n, sorted(touched)

    def _selected_run(self) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        run_key = self.selected_run_key
        if not run_key:
            run_key = self.state.get("current_key")
        if not run_key:
            return None, None
        runs = self.state.get("runs", {}) if isinstance(self.state.get("runs"), dict) else {}
        run = runs.get(run_key) if isinstance(runs.get(run_key), dict) else None
        return run_key, run

    def _export_selected_run(self):
        run_key, run = self._selected_run()
        if not run_key or not run:
            messagebox.showerror("Ошибка", "Нет выбранного заезда")
            return
        path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv")],
                                            initialfile=f"run_{run_key}.csv")
        if not path:
            return
        self._export_runs_to_csv(path, only_run_key=run_key)
        messagebox.showinfo("Готово", f"Сохранено: {path}")

    def _export_all_runs(self):
        if not self.state.get("runs"):
            messagebox.showerror("Ошибка", "Нет данных")
            return
        path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv")],
                                            initialfile="all_runs.csv")
        if not path:
            return
        self._export_runs_to_csv(path, only_run_key=None)
        messagebox.showinfo("Готово", f"Сохранено: {path}")

    def _export_runs_to_csv(self, path: str, only_run_key: Optional[str]):
        runs = self.state.get("runs", {}) if isinstance(self.state.get("runs"), dict) else {}
        keys = [only_run_key] if only_run_key else list(runs.keys())

        all_splits = set()
        for k in keys:
            run = runs.get(k) or {}
            for sid in self._run_split_ids(run):
                all_splits.add(str(sid))
        all_splits_sorted = sorted(all_splits, key=split_sort_key)

        cols = ["run", "category", "bib", "name", "country"]
        if self.show_distance_var.get():
            cols.append("distance_m")
        cols += [f"S{sid}" for sid in all_splits_sorted] + ["finish", "status"]

        with open(path, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f, delimiter=";")
            w.writerow(cols)
            for run_key in keys:
                run = runs.get(run_key) or {}
                cat = self.settings.run_categories.get(run_key, "")
                ath = run.get("athletes") or {}
                order = run.get("bib_order")
                if isinstance(order, list) and order:
                    bibs = [safe_int_str(b) for b in order if safe_int_str(b)]
                else:
                    bibs = [safe_int_str(b) for b in (ath.keys() if isinstance(ath, dict) else [])]
                bibs = [b for b in bibs if b in ath]

                for bib in bibs:
                    a = ath.get(bib)
                    if not isinstance(a, dict):
                        continue
                    name, country = self._effective_meta(run_key, a)
                    splits = a.get("splits") or {}
                    finish = a.get("finish")
                    status = a.get("status") or ""
                    row = [run_key, cat, bib, name, country]
                    if self.show_distance_var.get():
                        split_count = len(splits) if isinstance(splits, dict) else 0
                        row.append(split_count * DIST_PER_SPLIT_M if split_count > 0 else "")
                    for sid in all_splits_sorted:
                        v = ""
                        if isinstance(splits, dict):
                            v = fmt_time(splits.get(str(sid)))
                        row.append(v)
                    row.append(fmt_time(finish))
                    row.append(status)
                    w.writerow(row)

    def _save_snapshot_json(self):
        if not self.state:
            messagebox.showerror("Ошибка", "Нет данных")
            return
        path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON", "*.json")],
                                            initialfile=f"snapshot_{int(now_ts())}.json")
        if not path:
            return
        snap = {
            "ts_saved": now_ts(),
            "run_categories": self.settings.run_categories,
            "state": self.state,
        }
        Path(path).write_text(json.dumps(snap, ensure_ascii=False, indent=2), encoding="utf-8")
        messagebox.showinfo("Готово", f"Сохранено: {path}")

    def _load_snapshot_json(self):
        path = filedialog.askopenfilename(filetypes=[("JSON", "*.json"), ("All files", "*.*")])
        if not path:
            return
        try:
            raw = Path(path).read_text(encoding="utf-8")
            data = json.loads(raw)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать JSON:\n{e}")
            return

        state = None
        cats = None
        if isinstance(data, dict) and "state" in data:
            state = data.get("state")
            cats = data.get("run_categories")
        else:
            state = data

        if not isinstance(state, dict):
            messagebox.showerror("Ошибка", "В этом JSON нет поля 'state' или структура не распознана.")
            return

        self.state = state

        if isinstance(cats, dict):
            try:
                self.settings.run_categories = {str(k): str(v) for k, v in cats.items()}
            except Exception:
                pass

        try:
            self._refresh_views()
        except Exception:
            pass

        messagebox.showinfo("Готово", f"Загружено: {path}")


    def _copy_selected_run_text(self):
        run_key, run = self._selected_run()
        if not run_key or not run:
            return
        txt = self._build_run_text(run_key, run)
        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
        except Exception:
            pass
        self.export_text.delete("1.0", "end")
        self.export_text.insert("1.0", txt)

    def _build_run_text(self, run_key: str, run: Dict[str, Any]) -> str:
        cat = self.settings.run_categories.get(run_key, "")
        start = run.get("start_time") or ""
        lines = []
        lines.append(f"Заезд: {run_key}")
        if cat:
            lines.append(f"Категория: {cat}")
        if start:
            lines.append(f"Старт: {start}")
        lines.append("")
        lines.append("№;Имя;Стр/Гор;Отсечки;Финиш;Статус")

        split_ids = self._run_split_ids(run)
        ath = run.get("athletes") or {}
        order = run.get("bib_order")
        if isinstance(order, list) and order:
            bibs = [safe_int_str(b) for b in order if safe_int_str(b)]
        else:
            bibs = [safe_int_str(b) for b in (ath.keys() if isinstance(ath, dict) else [])]
        bibs = [b for b in bibs if b in ath]

        for bib in bibs:
            a = ath.get(bib)
            if not isinstance(a, dict):
                continue
            name, country = self._effective_meta(run_key, a)
            splits = a.get("splits") or {}
            sp = []
            for sid in split_ids:
                if isinstance(splits, dict) and str(sid) in splits:
                    sp.append(f"S{sid}:{fmt_time(splits.get(str(sid)))}")
            finish = fmt_time(a.get("finish"))
            status = str(a.get("status") or "")
            lines.append(f"{bib};{name};{country};{' '.join(sp)};{finish};{status}")
        return "\n".join(lines) + "\n"

    def _update_export_preview(self):
        run_key, run = self._selected_run()
        if not run_key or not run:
            txt = "Нет выбранного заезда.\n"
        else:
            txt = self._build_run_text(run_key, run)
        try:
            cur = self.export_text.get("1.0", "end")
            if cur != txt:
                self.export_text.delete("1.0", "end")
                self.export_text.insert("1.0", txt)
        except Exception:
            pass

    def _edit_categories(self):
        d = CategoryDialog(self, self.settings.categories)
        d.wait_window()
        if not d.result:
            return
        old = list(self.settings.categories)
        new = list(d.result)

        self.settings.categories = new
        self.run_cat_cb["values"] = new
        self.roster_cat_cb["values"] = new

        if self.category_var.get() not in new:
            self.category_var.set(new[0])
        if self.run_category_var.get() not in new and self.run_category_var.get().strip():
            self.run_category_var.set(new[0])

        for c in new:
            self.roster.ensure_category(c)

        removed = set(old) - set(new)
        if removed:
            for rk, cat in list(self.settings.run_categories.items()):
                if cat in removed:
                    self.settings.run_categories[rk] = ""

        self._render_roster()
        self._refresh_views()

    def _append_log(self, s: str):
        s = str(s).rstrip()
        if not s:
            return
        try:
            self.log.insert("end", s + "\n")
            lines = int(self.log.index("end-1c").split(".")[0])
            if lines > 1500:
                self.log.delete("1.0", f"{lines-1500}.0")
            self.log.see("end")
        except Exception:
            pass

    def _auto_save_tick(self):
        self._save_settings()
        self.after(600, self._auto_save_tick)



    def _runs_text_select_all(self, _evt=None):
        if not hasattr(self, "runs_text"):
            return "break"
        try:
            self.runs_text.tag_add("sel", "1.0", "end-1c")
            self.runs_text.mark_set("insert", "1.0")
            self.runs_text.see("1.0")
        except Exception:
            pass
        return "break"

    def _copy_runs_text_selection(self, _evt=None):
        if not hasattr(self, "runs_text"):
            return "break"
        try:
            txt = self.runs_text.selection_get()
        except Exception:
            txt = ""
        if not txt:
            return "break"
        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
        except Exception:
            pass
        return "break"

    def _copy_runs_text_all(self):
        if not hasattr(self, "runs_text"):
            return
        try:
            txt = self.runs_text.get("1.0", "end-1c")
        except Exception:
            txt = ""
        if not txt:
            return
        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
        except Exception:
            pass

    def _runs_text_context_menu(self, evt):
        m = tk.Menu(self, tearoff=0)
        m.add_command(label="Копировать", command=lambda: self._copy_runs_text_selection())
        m.add_command(label="Копировать всё", command=self._copy_runs_text_all)
        m.add_separator()
        m.add_command(label="Выделить всё", command=lambda: self._runs_text_select_all())
        try:
            m.tk_popup(evt.x_root, evt.y_root)
        finally:
            try:
                m.grab_release()
            except Exception:
                pass

    def _render_runs_table_text(self):
        if not hasattr(self, "runs_text"):
            return

        items = self._runs_in_display_order()

        seen_split_ids: set = set()
        for _rk, _run in items:
            for _sid in self._run_split_ids(_run):
                seen_split_ids.add(str(_sid))
        split_ids = sorted(seen_split_ids, key=split_sort_key)

        cols = ["Заезд", "Дисциплина", "Раунд", "Кат.", "Дист.,м", "Старт", "Место", "№", "Имя", "Стр/Гор"]
        cols += [f"S{sid}" for sid in split_ids] + ["Финиш", "Статус"]

        flt = self.run_filter_var.get().strip().lower()

        def _ath_hay(r: Dict[str, Any]) -> str:
            return f"{r.get('bib','')} {r.get('name','')} {r.get('country','')} {r.get('finish','')} {r.get('status','')}".lower()

        lines: List[str] = []
        lines.append("	".join(cols))

        for run_key, run in items:
            start = str(run.get("start_time") or "").strip()
            ev = str(run.get("event") or run.get("discipline") or run.get("name") or "").strip()
            rnd = str(run.get("round") or run.get("phase") or run.get("heat") or "").strip()
            cat = self.settings.run_categories.get(run_key, "") or ""

            dm = run_distance_m(run)
            dm_s = str(dm) if dm is not None else ""

            rows = self._protocol_sort_rows(run_key, run)

            place = 0
            for r in rows:
                st_u = (r.get("status") or "").strip().upper()
                if r.get("finish_f") is not None and st_u not in ("DNS",):
                    place += 1
                    r["place"] = str(place)
                else:
                    r["place"] = ""

            if flt:
                run_hay = f"{run_key} {ev} {rnd} {cat} {start} {dm_s}".lower()
                if flt in run_hay:
                    shown = rows
                else:
                    shown = [r for r in rows if flt in _ath_hay(r)]
                if not shown:
                    continue
            else:
                shown = rows

            for r in shown:
                spm = self._splits_to_map(r.get("splits"))
                sp_vals = [spm.get(str(sid), "") for sid in split_ids]

                row = [
                    str(run_key),
                    str(ev),
                    str(rnd),
                    str(cat),
                    str(dm_s),
                    str(start),
                    str(r.get("place", "") or ""),
                    str(r.get("bib", "") or ""),
                    str(r.get("name", "") or ""),
                    str(r.get("country", "") or ""),
                ]
                row += [str(x or "") for x in sp_vals]
                row += [str(r.get("finish", "") or ""), str(r.get("status", "") or "")]
                lines.append("\t".join(row))

        out = "\n".join(lines).rstrip() + "\n"

        try:
            # keep selection possible; block typing via <Key> binding
            self.runs_text.delete("1.0", "end")
            self.runs_text.insert("1.0", out)
            self.runs_text.mark_set("insert", "1.0")
        except Exception:
            pass

    def _save_settings(self):
        try:
            self.settings.host = self.host_var.get().strip() or self.settings.host
            try:
                self.settings.port = int(self.port_var.get().strip())
            except Exception:
                pass
            self.settings.auto_reconnect = bool(self.auto_reconnect_var.get())
            self.settings.override_server_names = bool(self.override_names_var.get())
            self.settings.show_distance = bool(self.show_distance_var.get())
            self.settings.secretary_name = self.secretary_var.get().strip()
            self.settings.chief_judge_name = self.chief_judge_var.get().strip()
            pm = (self.protocol_mode_var.get() or "time").strip().lower()
            if pm in ("time", "order"):
                self.settings.protocol_mode = pm
            self.settings.protocol_include_splits = bool(self.protocol_include_splits_var.get())

            self.settings.protocol_date = (self.protocol_date_var.get() or "").strip()
            self.settings.protocol_conditions = (self.protocol_conditions_var.get() or "").strip()
            self.settings_store.save(self.settings)
        except Exception:
            pass

    def _on_close(self):
        self._save_settings()
        try:
            self.disconnect()
        except Exception:
            pass
        self.destroy()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--host", default=None)
    ap.add_argument("--port", type=int, default=None)
    args = ap.parse_args()

    app = App()
    if args.host:
        app.host_var.set(args.host)
    if args.port:
        app.port_var.set(str(args.port))

    app.mainloop()


if __name__ == "__main__":
    main()
