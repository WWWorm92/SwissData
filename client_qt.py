import argparse
import copy
import concurrent.futures
import csv
import datetime
import json
import queue
import re
import socket
import threading
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from core_server import fmt_time, split_sort_key

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_ORIENT = None
    WD_TABLE_ALIGNMENT = None
    Cm = None
    Pt = None
    OxmlElement = None
    qn = None

try:
    from PySide6.QtCore import QTimer, Qt
    from PySide6.QtGui import QColor, QAction, QKeySequence
    from PySide6.QtWidgets import (
        QApplication,
        QCheckBox,
        QComboBox,
        QFileDialog,
        QFrame,
        QGridLayout,
        QHBoxLayout,
        QHeaderView,
        QLabel,
        QLineEdit,
        QMainWindow,
        QMessageBox,
        QPushButton,
        QPlainTextEdit,
        QSplitter,
        QTableWidget,
        QTableWidgetItem,
        QTreeWidget,
        QTreeWidgetItem,
        QAbstractItemView,
        QTabWidget,
        QVBoxLayout,
        QWidget,
    )
except Exception:
    print("PySide6 is required. Install with: py -3 -m pip install PySide6")
    raise


DIST_STEP_M = 125
SETTINGS_PATH = Path("client_qt_settings.json")
ROSTER_PATH = Path("client_qt_roster.json")
PROTOCOL_TYPES = [
    "Произвольно",
    "Гит 125 м с места",
    "Гит 250 м с места",
    "Гит 500 м с места",
    "Гит 1000 м с места",
    "Гонка преследования 2 км",
    "Гонка преследования 3 км",
    "Гонка преследования 4 км",
]


def safe_int_str(x: Any) -> str:
    s = str(x).strip()
    if not s:
        return ""
    if s.isdigit():
        return s
    m = re.search(r"(\d+)", s)
    return str(int(m.group(1))) if m else ""


def fmt_ru_time(v: Any) -> str:
    return str(fmt_time(v) or "").replace(".", ",")


def fmt_sec_ru(sec: Any) -> str:
    if sec is None:
        return ""
    try:
        f = float(sec)
    except Exception:
        return str(sec)
    return f"{f:.3f}".replace(".", ",")


def fmt_speed_kmh_ru(distance_m: Optional[int], sec: Any) -> str:
    if not distance_m:
        return ""
    try:
        t = float(sec)
    except Exception:
        return ""
    if t <= 0:
        return ""
    return f"{(float(distance_m) / t) * 3.6:.3f}".replace(".", ",")


def parse_distance_m(run: Dict[str, Any]) -> Optional[int]:
    for k in ("distance_m", "distance", "dist", "dist_m"):
        if k not in run:
            continue
        try:
            return int(float(run.get(k)))
        except Exception:
            pass
    return None


def run_key_sort(k: str) -> Tuple[int, int, str]:
    m = re.match(r"^\s*(\d+)\s*[-:]\s*(\d+)\s*$", str(k))
    if m:
        return (int(m.group(1)), int(m.group(2)), str(k))
    return (10**9, 10**9, str(k))


def checkpoints_count(run: Dict[str, Any]) -> int:
    split_ids = set()
    has_finish = False
    ath = run.get("athletes") or {}
    if isinstance(ath, dict):
        for _bib, a in ath.items():
            if not isinstance(a, dict):
                continue
            sp = a.get("splits")
            if isinstance(sp, dict):
                split_ids.update(str(k) for k in sp.keys())
            if a.get("finish") is not None:
                has_finish = True
    return len(split_ids) + (1 if has_finish else 0)


def infer_discipline(run: Dict[str, Any]) -> str:
    dm = None
    for k in ("distance_m", "distance", "dist", "dist_m"):
        if k in run:
            try:
                dm = int(float(run.get(k)))
                break
            except Exception:
                pass
    if dm is None:
        c = checkpoints_count(run)
        by_cp = {
            1: 125,
            2: 250,
            4: 500,
            8: 1000,
            16: 2000,
            24: 3000,
            32: 4000,
        }
        dm = by_cp.get(c)
    by_dm = {
        125: "Гит 125 м с места",
        250: "Гит 250 м с места",
        500: "Гит 500 м с места",
        1000: "Гит 1000 м с места",
        2000: "Гонка преследования 2 км",
        3000: "Гонка преследования 3 км",
        4000: "Гонка преследования 4 км",
    }
    return by_dm.get(dm, "Произвольно")


class NetThread(threading.Thread):
    def __init__(self, q: queue.Queue, stop_evt: threading.Event, host: str, port: int):
        super().__init__(daemon=True)
        self.q = q
        self.stop_evt = stop_evt
        self.host = host
        self.port = port
        self.sock: Optional[socket.socket] = None

    def run(self):
        buf = b""
        try:
            self.sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            self.sock.settimeout(2.0)
            self.sock.connect((self.host, self.port))
            self.q.put(("status", "connected"))
            self.sock.settimeout(0.5)

            while not self.stop_evt.is_set():
                try:
                    chunk = self.sock.recv(8192)
                    if chunk == b"":
                        raise RuntimeError("server disconnected")
                    buf += chunk
                    while b"\n" in buf:
                        line, buf = buf.split(b"\n", 1)
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            obj = json.loads(line.decode("utf-8", errors="ignore"))
                        except Exception:
                            continue
                        t = obj.get("type")
                        if t == "state" and isinstance(obj.get("state"), dict):
                            self.q.put(("state", obj["state"]))
                except socket.timeout:
                    continue
        except Exception as e:
            self.q.put(("err", str(e)))
        finally:
            try:
                if self.sock:
                    self.sock.close()
            except Exception:
                pass


def _local_ipv4s() -> List[str]:
    out: List[str] = []
    try:
        hn = socket.gethostname()
        for ip in socket.gethostbyname_ex(hn)[2]:
            if ip and "." in ip and not ip.startswith("127."):
                out.append(ip)
    except Exception:
        pass
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        if ip and "." in ip and not ip.startswith("127."):
            out.append(ip)
    except Exception:
        pass
    uniq: List[str] = []
    seen = set()
    for ip in out:
        if ip not in seen:
            seen.add(ip)
            uniq.append(ip)
    return uniq


def _scan_candidates(preferred_host: str) -> List[str]:
    hosts: List[str] = []
    seen = set()

    def add(h: str):
        hs = str(h or "").strip()
        if not hs or hs in seen:
            return
        seen.add(hs)
        hosts.append(hs)

    add(preferred_host)
    add("127.0.0.1")
    add("localhost")

    for ip in _local_ipv4s():
        parts = ip.split(".")
        if len(parts) != 4:
            continue
        pfx = ".".join(parts[:3])
        add(ip)
        for n in range(1, 255):
            cand = f"{pfx}.{n}"
            add(cand)
    return hosts


def _probe_host(host: str, port: int, timeout_sec: float = 0.2) -> bool:
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.settimeout(timeout_sec)
        s.connect((host, port))
        try:
            s.settimeout(0.15)
            _ = s.recv(32)
        except Exception:
            pass
        s.close()
        return True
    except Exception:
        return False


class ScanThread(threading.Thread):
    def __init__(self, q: queue.Queue, stop_evt: threading.Event, preferred_host: str, port: int):
        super().__init__(daemon=True)
        self.q = q
        self.stop_evt = stop_evt
        self.preferred_host = preferred_host
        self.port = port

    def run(self):
        hosts = _scan_candidates(self.preferred_host)
        self.q.put(("scan_status", f"Сканирование сети ({len(hosts)} адресов)..."))
        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=96) as ex:
                fut2host = {ex.submit(_probe_host, h, self.port, 0.22): h for h in hosts}
                for fut in concurrent.futures.as_completed(fut2host):
                    if self.stop_evt.is_set():
                        return
                    h = fut2host[fut]
                    ok = False
                    try:
                        ok = bool(fut.result())
                    except Exception:
                        ok = False
                    if ok:
                        self.q.put(("scan_found", (h, self.port)))
                        self.stop_evt.set()
                        return
        finally:
            if not self.stop_evt.is_set():
                self.q.put(("scan_none", None))


class ClientQtApp(QMainWindow):
    def __init__(self, host: str, port: int):
        super().__init__()
        self.setWindowTitle("Quantum Client (Qt)")
        self.resize(1620, 960)

        self.q = queue.Queue()
        self.stop_evt = threading.Event()
        self.net: Optional[NetThread] = None
        self.scan_stop_evt = threading.Event()
        self.scan_thr: Optional[ScanThread] = None
        self._connected = False
        self._last_connect_try = 0.0
        self.state: Dict[str, Any] = {}
        self.selected_run_key: Optional[str] = None
        self.run_categories: Dict[str, str] = {}
        self.roster_by_cat: Dict[str, Dict[str, Dict[str, str]]] = {}
        self.protocol_rejected: Dict[str, str] = {}
        self.protocol_suggestion: Optional[Dict[str, Any]] = None

        self._build_ui(host, port)
        self._apply_qss()
        self._load_settings()
        self._load_roster()
        self._refresh_roster_table()

        if self.auto_connect_cb.isChecked():
            if self.auto_scan_cb.isChecked():
                self.start_scan_connect()
            else:
                self.connect_net(silent=True)

        self.timer = QTimer(self)
        self.timer.timeout.connect(self._pump)
        self.timer.start(50)

    def _build_ui(self, host: str, port: int):
        root = QWidget(self)
        self.setCentralWidget(root)
        main = QVBoxLayout(root)
        main.setContentsMargins(18, 16, 18, 16)
        main.setSpacing(10)

        title = QLabel("Quantum Client")
        title.setObjectName("Title")
        sub = QLabel("Live results · categories · protocol export")
        sub.setObjectName("Sub")
       

        bar = QFrame(); bar.setObjectName("Card")
        bl = QHBoxLayout(bar)
        bl.setContentsMargins(12, 12, 12, 12)
        bl.setSpacing(10)

        self.host_ed = QLineEdit(host)
        self.host_ed.setMinimumWidth(190)
        self.port_ed = QLineEdit(str(port))
        self.port_ed.setMaximumWidth(100)
        self.cat_cb = QComboBox()
        self.cat_cb.addItems(["ALL"])
        self.cat_cb.currentTextChanged.connect(self._refresh_views)
        self.dist_cb = QCheckBox("Дистанция")
        self.dist_cb.setChecked(True)
        self.dist_cb.toggled.connect(self._refresh_athletes)
        self.override_cb = QCheckBox("Подменять имена из состава")
        self.override_cb.setChecked(True)
        self.override_cb.toggled.connect(self._refresh_views)
        self.auto_connect_cb = QCheckBox("Автоподключение")
        self.auto_connect_cb.setChecked(True)
        self.auto_scan_cb = QCheckBox("Автопоиск")
        self.auto_scan_cb.setChecked(True)

        b_connect = QPushButton("Подключить")
        b_connect.setObjectName("Accent")
        b_connect.clicked.connect(self.connect_net)
        b_disconnect = QPushButton("Отключить")
        b_disconnect.clicked.connect(self.disconnect_net)
        b_scan = QPushButton("Сканировать")
        b_scan.clicked.connect(self.start_scan_connect)

        for w in [QLabel("Host"), self.host_ed, QLabel("Port"), self.port_ed, QLabel("Категория"), self.cat_cb, self.dist_cb, self.override_cb, self.auto_connect_cb, self.auto_scan_cb, b_scan, b_connect, b_disconnect]:
            bl.addWidget(w)
        bl.addStretch(1)

        self.status_lbl = QLabel("Отключено")
        self.status_lbl.setObjectName("StatusBad")
        bl.addWidget(self.status_lbl)
        main.addWidget(bar)

        tabs = QTabWidget()
        main.addWidget(tabs, 1)

        tab_res = QWidget(); tab_roster = QWidget(); tab_protocol = QWidget(); tab_log = QWidget()
        tabs.addTab(tab_res, "Результаты")
        tabs.addTab(tab_roster, "Состав")
        tabs.addTab(tab_protocol, "Протокол")
        tabs.addTab(tab_log, "Лог")

        # Results
        rl = QVBoxLayout(tab_res)
        split = QSplitter(Qt.Horizontal)
        rl.addWidget(split)

        left = QFrame(); left.setObjectName("Card")
        right = QFrame(); right.setObjectName("Card")
        split.addWidget(left); split.addWidget(right)
        split.setSizes([460, 1100])

        ll = QVBoxLayout(left)
        ll.setContentsMargins(12, 12, 12, 12)
        ll.setSpacing(10)
        ll.addWidget(QLabel("Заезды"))

        runs_cat_row = QHBoxLayout()
        runs_cat_row.addWidget(QLabel("Категория"))
        self.runs_cat_cb = QComboBox()
        self.runs_cat_cb.addItems(["ALL"])
        self.runs_cat_cb.currentTextChanged.connect(self._refresh_views)
        runs_cat_row.addWidget(self.runs_cat_cb, 1)
        ll.addLayout(runs_cat_row)

        self.run_filter = QLineEdit()
        self.run_filter.setPlaceholderText("Фильтр по заездам")
        self.run_filter.textChanged.connect(self._refresh_runs)
        ll.addWidget(self.run_filter)

        self.runs_table = QTableWidget(0, 5)
        self.runs_table.setHorizontalHeaderLabels(["Заезд", "Кат.", "Старт", "Участн.", "Финиш"])
        self.runs_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.runs_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.runs_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.runs_table.setAlternatingRowColors(True)
        self.runs_table.verticalHeader().setVisible(False)
        self.runs_table.verticalHeader().setDefaultSectionSize(40)
        self.runs_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.runs_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.runs_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.runs_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.runs_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        self.runs_table.cellClicked.connect(self._on_run_selected)
        ll.addWidget(self.runs_table, 1)

        rr = QVBoxLayout(right)
        rr.setContentsMargins(12, 12, 12, 12)
        rr.setSpacing(10)
        top = QHBoxLayout()
        top.addWidget(QLabel("Участники"))
        top.addStretch(1)
        self.ath_filter = QLineEdit()
        self.ath_filter.setPlaceholderText("Фильтр участников")
        self.ath_filter.textChanged.connect(self._refresh_athletes)
        self.ath_filter.setMaximumWidth(240)
        top.addWidget(self.ath_filter)
        self.run_info = QLabel("—")
        self.run_info.setObjectName("Sub")
        top.addWidget(self.run_info)
        rr.addLayout(top)

        self.views_tabs = QTabWidget()
        rr.addWidget(self.views_tabs, 1)

        view_table = QWidget()
        vt_l = QVBoxLayout(view_table)
        vt_l.setContentsMargins(0, 0, 0, 0)
        self.ath_table = QTableWidget(0, 4)
        self.ath_table.setHorizontalHeaderLabels(["№", "Имя", "Финиш", "Статус"])
        self.ath_table.setSelectionBehavior(QTableWidget.SelectItems)
        self.ath_table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.ath_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.ath_table.setAlternatingRowColors(True)
        self.ath_table.verticalHeader().setVisible(False)
        self.ath_table.verticalHeader().setDefaultSectionSize(40)
        self.ath_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.ath_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.copy_cells_act = QAction(self)
        self.copy_cells_act.setShortcut(QKeySequence.Copy)
        self.copy_cells_act.triggered.connect(self._copy_selected_ath_cells)
        self.ath_table.addAction(self.copy_cells_act)
        vt_l.addWidget(self.ath_table)

        view_tree = QWidget()
        vtr_l = QVBoxLayout(view_tree)
        vtr_l.setContentsMargins(0, 0, 0, 0)
        self.ath_tree = QTreeWidget()
        self.ath_tree.setHeaderLabels(["Поле", "Значение"])
        self.ath_tree.setAlternatingRowColors(True)
        self.ath_tree.setColumnWidth(0, 260)
        vtr_l.addWidget(self.ath_tree)

        view_text = QWidget()
        vtx_l = QVBoxLayout(view_text)
        vtx_l.setContentsMargins(0, 0, 0, 0)
        self.ath_text = QPlainTextEdit()
        self.ath_text.setReadOnly(True)
        vtx_l.addWidget(self.ath_text)

        self.views_tabs.addTab(view_table, "Таблица")
        self.views_tabs.addTab(view_tree, "Дерево")
        self.views_tabs.addTab(view_text, "Текст")

        # Roster tab
        ro = QVBoxLayout(tab_roster)
        ro_card = QFrame(); ro_card.setObjectName("Card")
        ro_l = QVBoxLayout(ro_card)
        ro_l.setContentsMargins(12, 12, 12, 12)
        ro_l.setSpacing(10)

        ro_top = QHBoxLayout()
        ro_top.addWidget(QLabel("Категория"))
        self.roster_cat_cb = QComboBox()
        self.roster_cat_cb.addItems(["DEFAULT"])
        self.roster_cat_cb.currentTextChanged.connect(self._refresh_roster_table)
        ro_top.addWidget(self.roster_cat_cb, 1)
        b_import = QPushButton("Импорт Excel/CSV…")
        b_import.clicked.connect(self.import_roster)
        b_save_roster = QPushButton("Сохранить состав")
        b_save_roster.clicked.connect(self._save_roster)
        b_load_roster = QPushButton("Загрузить состав")
        b_load_roster.clicked.connect(self.load_roster_dialog)
        ro_top.addWidget(b_import)
        ro_top.addWidget(b_save_roster)
        ro_top.addWidget(b_load_roster)
        ro_l.addLayout(ro_top)

        self.roster_table = QTableWidget(0, 3)
        self.roster_table.setHorizontalHeaderLabels(["№", "Имя", "Страна/Регион"])
        self.roster_table.setAlternatingRowColors(True)
        self.roster_table.verticalHeader().setVisible(False)
        self.roster_table.verticalHeader().setDefaultSectionSize(38)
        self.roster_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.roster_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.roster_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.roster_table.itemChanged.connect(self._on_roster_item_changed)
        ro_l.addWidget(self.roster_table, 1)
        ro.addWidget(ro_card, 1)

        # Protocol
        pl = QVBoxLayout(tab_protocol)
        card = QFrame(); card.setObjectName("Card")
        gl = QGridLayout(card)
        gl.setContentsMargins(14, 14, 14, 14)
        gl.setHorizontalSpacing(10)
        gl.setVerticalSpacing(10)

        self.scope_cb = QComboBox()
        self.scope_cb.addItems(["Все заезды", "Текущий заезд", "Выбранные заезды", "По фильтру"])
        self.scope_cb.currentTextChanged.connect(self._refresh_protocol_hint)
        self.type_lbl = QLabel("Авто: —")
        self.protocol_type_cb = QComboBox()
        self.protocol_type_cb.addItems(PROTOCOL_TYPES)
        self.protocol_type_cb.currentTextChanged.connect(self._on_protocol_type_changed)
        self.event_ed = QLineEdit("")
        self.event_ed.setPlaceholderText("Дисциплина")
        self.event_ed.textChanged.connect(self._refresh_protocol_hint)
        self.round_ed = QLineEdit("")
        self.round_ed.setPlaceholderText("Раунд/заезд")
        self.round_ed.textChanged.connect(self._refresh_protocol_hint)
        self.date_ed = QLineEdit("")
        self.date_ed.setPlaceholderText("Дата (например: Воскресенье 8 февраля 2026 г.)")
        self.date_ed.textChanged.connect(self._refresh_protocol_hint)
        self.cond_ed = QLineEdit("")
        self.cond_ed.setPlaceholderText("Условия (например: t°C 20; P 991; vl.65%)")
        self.cond_ed.textChanged.connect(self._refresh_protocol_hint)
        self.secretary_ed = QLineEdit("")
        self.secretary_ed.setPlaceholderText("Секретарь")
        self.secretary_ed.textChanged.connect(self._refresh_protocol_hint)
        self.judge_ed = QLineEdit("")
        self.judge_ed.setPlaceholderText("Главный судья")
        self.judge_ed.textChanged.connect(self._refresh_protocol_hint)
        self.notes_ed = QLineEdit("")
        self.notes_ed.setPlaceholderText("Примечание")
        self.notes_ed.textChanged.connect(self._refresh_protocol_hint)
        self.sort_mode_cb = QComboBox()
        self.sort_mode_cb.addItems(["По времени", "По порядку"])
        self.sort_mode_cb.currentTextChanged.connect(self._refresh_protocol_hint)
        self.include_splits_cb = QCheckBox("Отсечки")
        self.include_splits_cb.toggled.connect(self._refresh_protocol_hint)
        self.suggest_lbl = QLabel("Подсказка: —")

        b_accept = QPushButton("Принять")
        b_reject = QPushButton("Отклонить")
        b_copy = QPushButton("Копировать")
        b_txt = QPushButton("TXT…")
        b_docx = QPushButton("DOCX…")
        b_docx_folder = QPushButton("Папка DOCX…")
        b_xlsx = QPushButton("XLSX…")
        b_accept.clicked.connect(self._apply_protocol_suggestion)
        b_reject.clicked.connect(self._reject_protocol_suggestion)
        b_copy.clicked.connect(self._copy_protocol_text)
        b_txt.clicked.connect(self.export_protocol_txt)
        b_docx.clicked.connect(self.export_protocol_docx)
        b_docx_folder.clicked.connect(self.export_protocol_docx_folder)
        b_xlsx.clicked.connect(self.export_protocol_xlsx)

        gl.addWidget(QLabel("Область"), 0, 0)
        gl.addWidget(self.scope_cb, 0, 1)
        gl.addWidget(QLabel("Тип"), 0, 2)
        gl.addWidget(self.protocol_type_cb, 0, 3)
        gl.addWidget(self.type_lbl, 0, 4)

        gl.addWidget(QLabel("Дисциплина"), 1, 0)
        gl.addWidget(self.event_ed, 1, 1, 1, 2)
        gl.addWidget(QLabel("Раунд"), 1, 3)
        gl.addWidget(self.round_ed, 1, 4)

        gl.addWidget(QLabel("Дата"), 2, 0)
        gl.addWidget(self.date_ed, 2, 1, 1, 2)
        gl.addWidget(QLabel("Условия"), 2, 3)
        gl.addWidget(self.cond_ed, 2, 4)

        gl.addWidget(QLabel("Секретарь"), 3, 0)
        gl.addWidget(self.secretary_ed, 3, 1)
        gl.addWidget(QLabel("Главный судья"), 3, 2)
        gl.addWidget(self.judge_ed, 3, 3, 1, 2)

        gl.addWidget(QLabel("Сортировка"), 4, 0)
        gl.addWidget(self.sort_mode_cb, 4, 1)
        gl.addWidget(self.include_splits_cb, 4, 2)
        gl.addWidget(self.notes_ed, 4, 3, 1, 2)

        gl.addWidget(self.suggest_lbl, 5, 0, 1, 3)
        gl.addWidget(b_accept, 5, 3)
        gl.addWidget(b_reject, 5, 4)

        gl.addWidget(b_copy, 6, 0)
        gl.addWidget(b_txt, 6, 1)
        gl.addWidget(b_docx, 6, 2)
        gl.addWidget(b_docx_folder, 6, 3)
        gl.addWidget(b_xlsx, 6, 4)
        self.protocol_preview = QPlainTextEdit()
        self.protocol_preview.setReadOnly(True)
        gl.addWidget(self.protocol_preview, 7, 0, 1, 5)
        pl.addWidget(card)
        pl.addStretch(1)

        # Log
        lg = QVBoxLayout(tab_log)
        self.log = QPlainTextEdit()
        self.log.setReadOnly(True)
        lg.addWidget(self.log)

    def _apply_qss(self):
        self.setStyleSheet(
            """
            QMainWindow, QWidget { background: #0b1118; color: #edf2f8; font-family: 'Segoe UI'; font-size: 13px; }
            #Title { font-size: 28px; font-weight: 800; }
            #Sub { color: #9ab0c8; }
            #Card { background: #13202c; border-radius: 14px; border: 1px solid #223549; }
            QLabel#StatusBad { background: #42242a; color: #ff9a9a; border-radius: 10px; padding: 8px 12px; font-weight: 700; }
            QLabel#StatusOk { background: #1f4132; color: #b2f0d0; border-radius: 10px; padding: 8px 12px; font-weight: 700; }
            QPushButton { background: #1f3144; color: #edf2f8; border: none; border-radius: 10px; padding: 9px 14px; font-weight: 600; }
            QPushButton:hover { background: #29465f; }
            QPushButton:pressed { background: #21374c; }
            QPushButton#Accent { background: #57c6ff; color: #091624; font-weight: 800; }
            QLineEdit, QPlainTextEdit, QTableWidget, QTreeWidget { background: #0f1a27; border: 1px solid #2a3f54; border-radius: 10px; }
            QComboBox { background: #0f1a27; border: 1px solid #2a3f54; border-radius: 10px; padding: 8px 10px; }
            QComboBox QAbstractItemView { background: #0f1a27; color: #edf2f8; selection-background-color: #2a4560; }
            QHeaderView::section { background: #1b2a3b; color: #edf2f8; border: none; padding: 8px; font-weight: 700; }
            QTableWidget, QTreeWidget { alternate-background-color: #122031; selection-background-color: #2a4560; }
            QTabWidget::pane { border: none; }
            QTabBar::tab { background: #13202c; color: #9eb2c8; border-radius: 10px; padding: 10px 16px; margin-right: 8px; }
            QTabBar::tab:selected { background: #1d3347; color: #edf2f8; }
            """
        )

    def _set_status(self, text: str, ok: bool):
        self.status_lbl.setText(text)
        self._connected = bool(ok)
        self.status_lbl.setObjectName("StatusOk" if ok else "StatusBad")
        self.status_lbl.style().unpolish(self.status_lbl)
        self.status_lbl.style().polish(self.status_lbl)

    def _load_settings(self):
        try:
            if not SETTINGS_PATH.exists():
                return
            d = json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
            if isinstance(d, dict):
                self.host_ed.setText(str(d.get("host") or self.host_ed.text()))
                self.port_ed.setText(str(d.get("port") or self.port_ed.text()))
                self.date_ed.setText(str(d.get("protocol_date") or ""))
                self.cond_ed.setText(str(d.get("protocol_conditions") or ""))
                self.event_ed.setText(str(d.get("protocol_event") or ""))
                self.round_ed.setText(str(d.get("protocol_round") or ""))
                self.secretary_ed.setText(str(d.get("secretary") or ""))
                self.judge_ed.setText(str(d.get("chief_judge") or ""))
                self.notes_ed.setText(str(d.get("protocol_notes") or ""))
                pt = str(d.get("protocol_type") or "Произвольно")
                if pt in PROTOCOL_TYPES:
                    self.protocol_type_cb.setCurrentText(pt)
                mode = str(d.get("protocol_sort_mode") or "По времени")
                if mode in ("По времени", "По порядку"):
                    self.sort_mode_cb.setCurrentText(mode)
                self.include_splits_cb.setChecked(bool(d.get("protocol_include_splits", False)))
                if bool(d.get("show_dist", True)) != self.dist_cb.isChecked():
                    self.dist_cb.setChecked(bool(d.get("show_dist", True)))
                self.override_cb.setChecked(bool(d.get("override_roster", True)))
                self.auto_connect_cb.setChecked(bool(d.get("auto_connect", True)))
                self.auto_scan_cb.setChecked(bool(d.get("auto_scan", True)))
        except Exception:
            pass

    def _save_settings(self):
        try:
            d = {
                "host": self.host_ed.text().strip(),
                "port": self.port_ed.text().strip(),
                "protocol_date": self.date_ed.text().strip(),
                "protocol_conditions": self.cond_ed.text().strip(),
                "protocol_event": self.event_ed.text().strip(),
                "protocol_round": self.round_ed.text().strip(),
                "secretary": self.secretary_ed.text().strip(),
                "chief_judge": self.judge_ed.text().strip(),
                "protocol_notes": self.notes_ed.text().strip(),
                "protocol_type": self.protocol_type_cb.currentText().strip(),
                "protocol_sort_mode": self.sort_mode_cb.currentText().strip(),
                "protocol_include_splits": bool(self.include_splits_cb.isChecked()),
                "override_roster": bool(self.override_cb.isChecked()),
                "show_dist": bool(self.dist_cb.isChecked()),
                "auto_connect": bool(self.auto_connect_cb.isChecked()),
                "auto_scan": bool(self.auto_scan_cb.isChecked()),
            }
            SETTINGS_PATH.write_text(json.dumps(d, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass

    def _load_roster(self):
        try:
            if ROSTER_PATH.exists():
                d = json.loads(ROSTER_PATH.read_text(encoding="utf-8"))
                if isinstance(d, dict):
                    self.roster_by_cat = d
        except Exception:
            self.roster_by_cat = {}

    def _save_roster(self, silent: bool = False):
        try:
            ROSTER_PATH.write_text(json.dumps(self.roster_by_cat, ensure_ascii=False, indent=2), encoding="utf-8")
            if not silent:
                QMessageBox.information(self, "Готово", f"Состав сохранен: {ROSTER_PATH}")
        except Exception as e:
            if not silent:
                QMessageBox.critical(self, "Ошибка", str(e))

    def load_roster_dialog(self):
        p, _ = QFileDialog.getOpenFileName(self, "Загрузить состав", "", "JSON (*.json);;All (*.*)")
        if not p:
            return
        try:
            d = json.loads(Path(p).read_text(encoding="utf-8"))
            if not isinstance(d, dict):
                raise ValueError("Неверный формат JSON")
            self.roster_by_cat = d
            self._refresh_categories()
            self._refresh_roster_table()
            QMessageBox.information(self, "Готово", f"Загружено: {p}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def import_roster(self):
        p, _ = QFileDialog.getOpenFileName(self, "Импорт состава", "", "Excel/CSV (*.xlsx *.csv);;All (*.*)")
        if not p:
            return
        cat = self.roster_cat_cb.currentText().strip() or "DEFAULT"
        try:
            dst = self.roster_by_cat.setdefault(cat, {})
            if p.lower().endswith(".xlsx"):
                if load_workbook is None:
                    raise RuntimeError("openpyxl не установлен")
                wb = load_workbook(p, data_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(min_row=1, values_only=True))
                if not rows:
                    return
                header = [str(x or "").strip().lower() for x in rows[0]]
                has_header = any(h in ("bib", "номер", "№", "name", "имя", "country", "страна") for h in header)
                start = 1 if has_header else 0

                def idx(names: List[str], default: Optional[int]) -> Optional[int]:
                    if has_header:
                        for i, h in enumerate(header):
                            if h in names:
                                return i
                    return default

                i_bib = idx(["bib", "номер", "№"], 0)
                i_name = idx(["name", "имя", "фио"], 1)
                i_country = idx(["country", "страна", "город"], 2)
                i_org = idx(["org", "organization", "организация", "школа", "клуб", "команда"], None)
                i_dob = idx(["dob", "birthdate", "date_of_birth", "дата рождения", "д.р.", "др"], None)
                i_rank = idx(["rank", "разряд", "спорт разряд", "спортразряд"], None)
                i_region = idx(["region", "регион", "область", "город/регион"], None)

                for r in rows[start:]:
                    if not r:
                        continue
                    bib = safe_int_str(r[i_bib] if i_bib is not None and i_bib < len(r) else "")
                    if not bib or bib == "0":
                        continue
                    rec = dst.setdefault(bib, {})
                    rec["name"] = str(r[i_name] or "").strip() if i_name is not None and i_name < len(r) else str(rec.get("name") or "")
                    rec["country"] = (str(r[i_country] or "").strip().upper() if i_country is not None and i_country < len(r) else str(rec.get("country") or "").upper())
                    for key, i_col in (("org", i_org), ("dob", i_dob), ("rank", i_rank), ("region", i_region)):
                        if i_col is not None and i_col < len(r):
                            rec[key] = str(r[i_col] or "").strip()
            else:
                with open(p, "r", encoding="utf-8-sig", newline="") as f:
                    text = f.read()
                lines = text.splitlines()
                if not lines:
                    return
                delim = ";" if lines[0].count(";") >= lines[0].count(",") else ","
                rows = list(csv.reader(lines, delimiter=delim))
                header = [str(x or "").strip().lower() for x in rows[0]]
                has_header = any(h in ("bib", "номер", "№", "name", "имя", "country", "страна") for h in header)
                start = 1 if has_header else 0

                def idx(names: List[str], default: Optional[int]) -> Optional[int]:
                    if has_header:
                        for i, h in enumerate(header):
                            if h in names:
                                return i
                    return default

                i_bib = idx(["bib", "номер", "№"], 0)
                i_name = idx(["name", "имя", "фио"], 1)
                i_country = idx(["country", "страна", "город"], 2)
                i_org = idx(["org", "organization", "организация", "школа", "клуб", "команда"], None)
                i_dob = idx(["dob", "birthdate", "date_of_birth", "дата рождения", "д.р.", "др"], None)
                i_rank = idx(["rank", "разряд", "спорт разряд", "спортразряд"], None)
                i_region = idx(["region", "регион", "область", "город/регион"], None)

                for row in rows[start:]:
                    if not row:
                        continue
                    bib = safe_int_str(row[i_bib] if i_bib is not None and i_bib < len(row) else "")
                    if not bib or bib == "0":
                        continue
                    rec = dst.setdefault(bib, {})
                    rec["name"] = str(row[i_name] or "").strip() if i_name is not None and i_name < len(row) else str(rec.get("name") or "")
                    rec["country"] = (str(row[i_country] or "").strip().upper() if i_country is not None and i_country < len(row) else str(rec.get("country") or "").upper())
                    for key, i_col in (("org", i_org), ("dob", i_dob), ("rank", i_rank), ("region", i_region)):
                        if i_col is not None and i_col < len(row):
                            rec[key] = str(row[i_col] or "").strip()
            self._refresh_roster_table()
            self._save_roster(silent=True)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def _refresh_roster_table(self):
        cat = self.roster_cat_cb.currentText().strip() or "DEFAULT"
        data = self.roster_by_cat.get(cat, {}) if isinstance(self.roster_by_cat, dict) else {}
        self.roster_table.blockSignals(True)
        self.roster_table.setRowCount(0)
        for bib in sorted(data.keys(), key=lambda x: int(x) if str(x).isdigit() else 10**9):
            rec = data.get(bib) or {}
            r = self.roster_table.rowCount()
            self.roster_table.insertRow(r)
            self.roster_table.setItem(r, 0, QTableWidgetItem(str(bib)))
            self.roster_table.setItem(r, 1, QTableWidgetItem(str(rec.get("name") or "")))
            self.roster_table.setItem(r, 2, QTableWidgetItem(str(rec.get("country") or "")))
        self.roster_table.blockSignals(False)

    def _on_roster_item_changed(self, item: QTableWidgetItem):
        try:
            r = item.row()
            bib_it = self.roster_table.item(r, 0)
            if not bib_it:
                return
            m = re.search(r"(\d+)", str(bib_it.text() or ""))
            if not m:
                return
            bib = str(int(m.group(1)))
            cat = self.roster_cat_cb.currentText().strip() or "DEFAULT"
            rec = self.roster_by_cat.setdefault(cat, {}).setdefault(bib, {"name": "", "country": ""})
            rec["name"] = str((self.roster_table.item(r, 1).text() if self.roster_table.item(r, 1) else "") or "").strip()
            rec["country"] = str((self.roster_table.item(r, 2).text() if self.roster_table.item(r, 2) else "") or "").strip().upper()
        except Exception:
            pass

    def _effective_meta(self, run_key: str, bib: str, a: Dict[str, Any]) -> Tuple[str, str]:
        srv_name = str(a.get("name") or "")
        srv_country = str(a.get("country") or "")
        cat = self.run_categories.get(run_key, "")
        if not cat:
            return srv_name, srv_country
        rec = (self.roster_by_cat.get(cat) or {}).get(str(bib)) if isinstance(self.roster_by_cat, dict) else None
        if not isinstance(rec, dict):
            return srv_name, srv_country
        use_override = bool(self.override_cb.isChecked()) if hasattr(self, "override_cb") else True
        nm = (str(rec.get("name") or "").strip() if use_override or not srv_name else "") or srv_name
        cc = (str(rec.get("country") or "").strip().upper() if use_override or not srv_country else "") or srv_country
        return nm, cc

    def _effective_meta_full(self, run_key: str, bib: str, a: Dict[str, Any]) -> Dict[str, str]:
        nm, cc = self._effective_meta(run_key, bib, a)
        out = {
            "name": nm,
            "country": cc,
            "org": str(a.get("org") or a.get("club") or a.get("team") or a.get("organization") or ""),
            "dob": str(a.get("dob") or a.get("birthdate") or a.get("birth_date") or a.get("date_of_birth") or ""),
            "rank": str(a.get("rank") or a.get("category") or a.get("class") or ""),
            "region": str(a.get("region") or a.get("city") or a.get("place") or ""),
        }
        cat = self.run_categories.get(run_key, "")
        if not cat:
            return out
        rec = (self.roster_by_cat.get(cat) or {}).get(str(bib)) if isinstance(self.roster_by_cat, dict) else None
        if not isinstance(rec, dict):
            return out
        use_override = bool(self.override_cb.isChecked()) if hasattr(self, "override_cb") else True

        def pick(k: str, default: str) -> str:
            v = str(rec.get(k) or "").strip()
            return v or default if (use_override or not default) else default

        out["name"] = pick("name", out["name"])
        out["country"] = pick("country", out["country"]).upper()
        for k in ("org", "dob", "rank", "region"):
            out[k] = pick(k, out.get(k, ""))
        return out

    def _append_log(self, text: str):
        if not text:
            return
        self.log.appendPlainText(str(text))

    def connect_net(self, silent: bool = False):
        if self.net and self.net.is_alive():
            return
        host = self.host_ed.text().strip() or "127.0.0.1"
        try:
            port = int(self.port_ed.text().strip())
        except Exception:
            if not silent:
                QMessageBox.critical(self, "Ошибка", "Неверный порт")
            else:
                self._append_log("ERROR: неверный порт")
            return
        self._last_connect_try = time.monotonic()
        self.stop_evt.clear()
        self.net = NetThread(self.q, self.stop_evt, host, port)
        self.net.start()

    def disconnect_net(self):
        self.stop_evt.set()
        self.scan_stop_evt.set()
        self._set_status("Отключено", False)

    def start_scan_connect(self):
        if self.scan_thr and self.scan_thr.is_alive():
            return
        if self.net and self.net.is_alive():
            return
        try:
            port = int(self.port_ed.text().strip())
        except Exception:
            QMessageBox.critical(self, "Ошибка", "Неверный порт")
            return
        self._last_connect_try = time.monotonic()
        self.scan_stop_evt.clear()
        preferred = self.host_ed.text().strip() or "127.0.0.1"
        self.scan_thr = ScanThread(self.q, self.scan_stop_evt, preferred, port)
        self.scan_thr.start()

    def _runs_items(self) -> List[Tuple[str, Dict[str, Any]]]:
        runs = self.state.get("runs") if isinstance(self.state.get("runs"), dict) else {}
        items = [(str(k), v) for k, v in runs.items() if isinstance(v, dict)]
        items.sort(key=lambda kv: run_key_sort(kv[0]))
        return items

    def _visible_run_items(self) -> List[Tuple[str, Dict[str, Any]]]:
        items = self._runs_items()
        flt = self.run_filter.text().strip().lower()
        cat = self.runs_cat_cb.currentText().strip() or "ALL"
        out: List[Tuple[str, Dict[str, Any]]] = []
        for run_key, run in items:
            run_cat = str(run.get("category") or "")
            if cat != "ALL" and run_cat != cat:
                continue
            hay = f"{run_key} {run_cat} {run.get('start_time') or ''} {len((run.get('athletes') or {}))}".lower()
            if flt and flt not in hay:
                continue
            out.append((run_key, run))
        return out

    def _selected_run_keys(self) -> List[str]:
        keys: List[str] = []
        try:
            rows = sorted({idx.row() for idx in self.runs_table.selectionModel().selectedRows()})
        except Exception:
            rows = []
        for r in rows:
            it = self.runs_table.item(r, 0)
            if it:
                k = it.text().strip()
                if k:
                    keys.append(k)
        return keys

    def _resolve_current_run_key(self, visible_items: Optional[List[Tuple[str, Dict[str, Any]]]] = None) -> Optional[str]:
        vis = visible_items if visible_items is not None else self._visible_run_items()
        if not vis:
            return None
        visible_keys = {k for k, _r in vis}
        if self.selected_run_key and self.selected_run_key in visible_keys:
            return self.selected_run_key
        return vis[0][0]

    def _refresh_categories(self):
        raw = self.state.get("categories") if isinstance(self.state, dict) else None
        cats = ["ALL"]
        roster_cats = sorted(list(self.roster_by_cat.keys())) if isinstance(self.roster_by_cat, dict) else []
        seen = set(cats)
        for rc in roster_cats:
            if rc and rc not in seen:
                seen.add(rc)
                cats.append(rc)
        if isinstance(raw, list):
            for x in raw:
                s = str(x or "").strip()
                if s and s not in seen:
                    seen.add(s)
                    cats.append(s)
        for _rk, rc in self.run_categories.items():
            if rc and rc not in seen:
                seen.add(rc)
                cats.append(rc)
        cur_runs_cat = self.runs_cat_cb.currentText() or "ALL"
        self.runs_cat_cb.blockSignals(True)
        self.runs_cat_cb.clear()
        self.runs_cat_cb.addItems(cats)
        if cur_runs_cat in cats:
            self.runs_cat_cb.setCurrentText(cur_runs_cat)
        else:
            self.runs_cat_cb.setCurrentText("ALL")
        self.runs_cat_cb.blockSignals(False)

        roster_values = [c for c in cats if c != "ALL"] or ["DEFAULT"]
        cur_roster = self.roster_cat_cb.currentText().strip() or "DEFAULT"
        self.roster_cat_cb.blockSignals(True)
        self.roster_cat_cb.clear()
        self.roster_cat_cb.addItems(roster_values)
        if cur_roster in roster_values:
            self.roster_cat_cb.setCurrentText(cur_roster)
        else:
            self.roster_cat_cb.setCurrentText(roster_values[0])
        self.roster_cat_cb.blockSignals(False)

    def _refresh_runs(self):
        visible = self._visible_run_items()
        current = self._resolve_current_run_key(visible)
        self.selected_run_key = current
        self.runs_table.setRowCount(0)
        for run_key, run in visible:
            run_cat = str(run.get("category") or "")
            r = self.runs_table.rowCount()
            self.runs_table.insertRow(r)
            vals = [run_key, run_cat, str(run.get("start_time") or ""), str(len(run.get("athletes") or {})), str(sum(1 for _b, a in (run.get("athletes") or {}).items() if isinstance(a, dict) and a.get("finish") is not None))]
            for cidx, v in enumerate(vals):
                it = QTableWidgetItem(v)
                if cidx != 1:
                    it.setTextAlignment(Qt.AlignCenter)
                self.runs_table.setItem(r, cidx, it)
            if run_key == current:
                self.runs_table.selectRow(r)

    def _on_run_selected(self, row: int, _col: int):
        it = self.runs_table.item(row, 0)
        if not it:
            return
        self.selected_run_key = it.text().strip()
        self._refresh_protocol_hint()

    def _refresh_athletes(self):
        visible_runs = self._visible_run_items()
        if not visible_runs:
            self.run_info.setText("—")
            self.ath_table.setRowCount(0)
            base = ["№", "Имя"]
            base.insert(0, "Заезд")
            if self.dist_cb.isChecked():
                base.append("Дист.")
            base += ["Финиш", "Статус"]
            self.ath_table.setColumnCount(len(base))
            self.ath_table.setHorizontalHeaderLabels(base)
            return

        run_split_ids: Dict[str, List[str]] = {}
        max_splits = 0
        for rk, run in visible_runs:
            sids = set()
            for _bib, a in (run.get("athletes") or {}).items():
                sp = a.get("splits") if isinstance(a, dict) else None
                if isinstance(sp, dict):
                    sids.update(str(k) for k in sp.keys())
            sorted_ids = sorted(sids, key=split_sort_key)
            run_split_ids[rk] = sorted_ids
            if len(sorted_ids) > max_splits:
                max_splits = len(sorted_ids)

        base_prefix = ["Заезд", "№", "Имя"]
        if self.dist_cb.isChecked():
            base_prefix.append("Дист.")
        headers = base_prefix + [f"S{i + 1}" for i in range(max_splits)] + ["Финиш", "Статус"]
        self.ath_table.setColumnCount(len(headers))
        self.ath_table.setHorizontalHeaderLabels(headers)

        self.ath_table.setRowCount(0)
        self.ath_tree.clear()
        self.ath_text.setPlainText("")
        text_lines: List[str] = ["\t".join(headers)]

        ath_filter = self.ath_filter.text().strip().lower()
        shown_total = 0
        for run_key, run in visible_runs:
            local_splits = run_split_ids.get(run_key, [])
            split_start = len(base_prefix)
            finish_col = split_start + len(local_splits)
            even_cols: List[int] = []
            for i, sid in enumerate(local_splits):
                try:
                    if int(str(sid)) % 2 == 0:
                        even_cols.append(split_start + i)
                except Exception:
                    pass

            ath = run.get("athletes") if isinstance(run.get("athletes"), dict) else {}
            order = run.get("bib_order") if isinstance(run.get("bib_order"), list) else list(ath.keys())

            sep_row = self.ath_table.rowCount()
            self.ath_table.insertRow(sep_row)
            sep_text = f"Заезд {run_key}   Кат.: {run.get('category') or '—'}"
            sep_item = QTableWidgetItem(sep_text)
            sep_item.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
            sep_item.setBackground(QColor("#1d2f43"))
            sep_item.setForeground(QColor("#d9ecff"))
            sep_item.setFlags(Qt.ItemIsEnabled)
            self.ath_table.setItem(sep_row, 0, sep_item)
            try:
                self.ath_table.setSpan(sep_row, 0, 1, len(headers))
            except Exception:
                pass

            hdr_row = self.ath_table.rowCount()
            self.ath_table.insertRow(hdr_row)
            local_headers = base_prefix + [f"S{sid}" for sid in local_splits] + ["Финиш", "Статус"]
            for cidx, txt in enumerate(local_headers):
                hi = QTableWidgetItem(txt)
                hi.setTextAlignment(Qt.AlignCenter)
                hi.setFlags(Qt.ItemIsEnabled)
                hi.setBackground(QColor("#22384d"))
                hi.setForeground(QColor("#d9ecff"))
                if cidx in even_cols or cidx == finish_col:
                    hi.setBackground(QColor("#1D4031"))
                    hi.setForeground(QColor("#aaf0c8"))
                self.ath_table.setItem(hdr_row, cidx, hi)

            run_node = QTreeWidgetItem([f"Заезд {run_key}", f"Кат.: {run.get('category') or '—'}"])
            self.ath_tree.addTopLevelItem(run_node)

            for bib in order:
                b = str(bib)
                a = ath.get(b)
                if not isinstance(a, dict):
                    continue
                nm, cc = self._effective_meta(run_key, b, a)
                hay = f"{run_key} {b} {nm} {cc} {a.get('status') or ''}".lower()
                if ath_filter and ath_filter not in hay:
                    continue

                shown_total += 1
                r = self.ath_table.rowCount()
                self.ath_table.insertRow(r)
                vals = [run_key, b, nm]
                if self.dist_cb.isChecked():
                    sp = a.get("splits") if isinstance(a.get("splits"), dict) else {}
                    cp = len(sp) + (1 if a.get("finish") is not None else 0)
                    vals.append(f"{cp * DIST_STEP_M}м" if cp else "")

                sp_map = a.get("splits") if isinstance(a.get("splits"), dict) else {}
                vals += [fmt_ru_time(sp_map.get(str(sid))) for sid in local_splits]
                vals += [fmt_ru_time(a.get("finish")), str(a.get("status") or "")]

                for cidx, v in enumerate(vals):
                    it = QTableWidgetItem(v)
                    if cidx != 2:
                        it.setTextAlignment(Qt.AlignCenter)
                    if cidx in even_cols:
                        it.setBackground(QColor("#1D4031"))
                    if cidx == finish_col:
                        it.setBackground(QColor("#1D4031"))
                    self.ath_table.setItem(r, cidx, it)

                n = QTreeWidgetItem([f"{b}  {nm}", f"{fmt_ru_time(a.get('finish'))}  {a.get('status') or ''}"])
                run_node.addChild(n)
                for sid in local_splits:
                    spv = fmt_ru_time(sp_map.get(str(sid)))
                    child = QTreeWidgetItem([f"S{sid}", spv])
                    try:
                        if int(str(sid)) % 2 == 0:
                            child.setBackground(0, QColor("#3b2f1a"))
                            child.setBackground(1, QColor("#3b2f1a"))
                    except Exception:
                        pass
                    n.addChild(child)

                text_lines.append("\t".join(str(x or "") for x in vals))

        self.ath_table.resizeColumnsToContents()
        for i in range(self.ath_tree.topLevelItemCount()):
            self.ath_tree.topLevelItem(i).setExpanded(True)
        self.ath_text.setPlainText("\n".join(text_lines))
        cat_text = self.runs_cat_cb.currentText().strip() or "ALL"
        self.run_info.setText(f"Кат.: {cat_text}   заездов: {len(visible_runs)}   показано участников: {shown_total}")

    def _refresh_views(self):
        self._refresh_runs()
        self._refresh_athletes()
        self._refresh_protocol_hint()

    def _refresh_protocol_hint(self):
        items = self._protocol_items()
        if not items:
            self.type_lbl.setText("Авто: —")
            self.suggest_lbl.setText("Подсказка: —")
            self.protocol_preview.setPlainText("")
            return
        run_key, run = items[0]
        self.type_lbl.setText(f"Авто: {infer_discipline(run)}")
        self._update_protocol_suggestion(run_key, run)
        self.protocol_preview.setPlainText(self._build_all_protocol_text(items))

    def _on_protocol_type_changed(self):
        t = self.protocol_type_cb.currentText().strip()
        if t and t != "Произвольно" and not self.event_ed.text().strip():
            self.event_ed.setText(t)
        self._refresh_protocol_hint()

    def _protocol_items(self) -> List[Tuple[str, Dict[str, Any]]]:
        items = self._visible_run_items()
        scope = self.scope_cb.currentIndex()
        if scope == 1:
            current = self._resolve_current_run_key(items)
            items = [(k, r) for k, r in items if current and k == current]
        elif scope == 2:
            sel = set(self._selected_run_keys())
            items = [(k, r) for k, r in items if k in sel]
        elif scope == 3:
            items = list(self._visible_run_items())
        return items

    def _run_checkpoint_count(self, run: Dict[str, Any]) -> int:
        split_ids = set()
        has_finish = False
        ath = run.get("athletes") if isinstance(run.get("athletes"), dict) else {}
        for _bib, a in ath.items():
            if not isinstance(a, dict):
                continue
            sp = a.get("splits")
            if isinstance(sp, dict):
                split_ids.update(str(k) for k in sp.keys())
            if a.get("finish") is not None:
                has_finish = True
        if not split_ids and has_finish:
            return 1
        return len(split_ids) + (1 if has_finish else 0)

    def _detect_protocol_discipline(self, run_key: str, run: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        dm = parse_distance_m(run)
        by_dm = {
            125: "Гит 125 м с места",
            250: "Гит 250 м с места",
            500: "Гит 500 м с места",
            1000: "Гит 1000 м с места",
            2000: "Гонка преследования 2 км",
            3000: "Гонка преследования 3 км",
            4000: "Гонка преследования 4 км",
        }
        if dm in by_dm:
            return {"run_key": run_key, "type": by_dm[dm], "event": by_dm[dm], "checkpoints": self._run_checkpoint_count(run), "signature": f"d:{dm}", "reason": "по дистанции"}
        cp = self._run_checkpoint_count(run)
        by_cp = {
            1: ("Гит 125 м с места", 125),
            2: ("Гит 250 м с места", 250),
            4: ("Гит 500 м с места", 500),
            8: ("Гит 1000 м с места", 1000),
            16: ("Гонка преследования 2 км", 2000),
            24: ("Гонка преследования 3 км", 3000),
            32: ("Гонка преследования 4 км", 4000),
        }
        if cp in by_cp:
            typ, _dm = by_cp[cp]
            return {"run_key": run_key, "type": typ, "event": typ, "checkpoints": cp, "signature": f"c:{cp}", "reason": "по числу отсечек"}
        return None

    def _update_protocol_suggestion(self, run_key: str, run: Dict[str, Any]):
        sug = self._detect_protocol_discipline(run_key, run)
        if not sug:
            self.protocol_suggestion = None
            self.suggest_lbl.setText("Подсказка: не удалось определить дисциплину")
            return
        sig = str(sug.get("signature") or "")
        if self.protocol_rejected.get(str(run_key)) == sig:
            self.protocol_suggestion = None
            self.suggest_lbl.setText("Подсказка: отклонена")
            return
        self.protocol_suggestion = sug
        self.suggest_lbl.setText(f"Подсказка: {sug.get('type')} ({sug.get('reason')})")

    def _apply_protocol_suggestion(self):
        sug = self.protocol_suggestion
        if not sug:
            return
        t = str(sug.get("type") or "").strip()
        if t and t in PROTOCOL_TYPES:
            self.protocol_type_cb.setCurrentText(t)
        ev = str(sug.get("event") or t)
        if ev:
            self.event_ed.setText(ev)
        self._refresh_protocol_hint()

    def _reject_protocol_suggestion(self):
        sug = self.protocol_suggestion
        if not sug:
            return
        rk = str(sug.get("run_key") or "")
        if rk:
            self.protocol_rejected[rk] = str(sug.get("signature") or "")
        self.protocol_suggestion = None
        self.suggest_lbl.setText("Подсказка: отклонена")

    def _effective_protocol_type_for_run(self, run_key: str, run: Dict[str, Any]) -> str:
        t = self.protocol_type_cb.currentText().strip()
        if t and t != "Произвольно":
            return t
        sug = self._detect_protocol_discipline(run_key, run)
        if sug:
            return str(sug.get("type") or "")
        return t

    def _protocol_sort_rows(self, run_key: str, run: Dict[str, Any]) -> List[Dict[str, Any]]:
        ath = run.get("athletes") if isinstance(run.get("athletes"), dict) else {}
        order = run.get("bib_order") if isinstance(run.get("bib_order"), list) else list(ath.keys())
        rows = []
        for idx, bib in enumerate(order):
            b = str(bib)
            a = ath.get(b)
            if not isinstance(a, dict):
                continue
            nm, cc = self._effective_meta(run_key, b, a)
            finish = a.get("finish")
            try:
                ff = float(finish) if finish is not None else None
            except Exception:
                ff = None
            rows.append({
                "idx": idx,
                "bib": b,
                "name": nm,
                "country": cc,
                "finish_f": ff,
                "finish": fmt_ru_time(finish),
                "status": str(a.get("status") or ""),
                "splits": a.get("splits") if isinstance(a.get("splits"), dict) else {},
            })
        if self.sort_mode_cb.currentIndex() == 1:
            rows.sort(key=lambda r: (r["finish_f"] is None, r["idx"]))
        else:
            rows.sort(key=lambda r: (r["finish_f"] is None, r["finish_f"] if r["finish_f"] is not None else 0.0, r["idx"]))
        place = 0
        for r in rows:
            st_u = (r.get("status") or "").strip().upper()
            if r.get("finish_f") is not None and st_u != "DNS":
                place += 1
                r["place"] = str(place)
            else:
                r["place"] = ""
        return rows

    def _protocol_groups(self, items: Optional[List[Tuple[str, Dict[str, Any]]]] = None) -> List[Tuple[Tuple[str, str], List[Tuple[str, Dict[str, Any]]]]]:
        use = items if items is not None else self._protocol_items()
        grouped: Dict[Tuple[str, str], List[Tuple[str, Dict[str, Any]]]] = {}
        order: List[Tuple[str, str]] = []
        for run_key, run in use:
            disc = self._effective_protocol_type_for_run(run_key, run) or infer_discipline(run) or "Произвольно"
            cat = str(run.get("category") or self.run_categories.get(run_key) or "").strip() or "Без категории"
            gk = (disc, cat)
            if gk not in grouped:
                grouped[gk] = []
                order.append(gk)
            grouped[gk].append((run_key, run))
        return [(gk, grouped[gk]) for gk in order]

    def _protocol_group_rows(self, group_runs: List[Tuple[str, Dict[str, Any]]]) -> Tuple[List[Dict[str, Any]], List[str]]:
        rows: List[Dict[str, Any]] = []
        split_ids_set = set()
        idx_global = 0
        for run_key, run in group_runs:
            ath = run.get("athletes") if isinstance(run.get("athletes"), dict) else {}
            order = run.get("bib_order") if isinstance(run.get("bib_order"), list) else list(ath.keys())
            for bib in order:
                b = str(bib)
                a = ath.get(b)
                if not isinstance(a, dict):
                    continue
                sp = a.get("splits") if isinstance(a.get("splits"), dict) else {}
                split_ids_set.update(str(k) for k in sp.keys())
                finish = a.get("finish")
                try:
                    ff = float(finish) if finish is not None else None
                except Exception:
                    ff = None
                meta = self._effective_meta_full(run_key, b, a)
                rows.append({
                    "idx": idx_global,
                    "run_key": run_key,
                    "bib": b,
                    "name": meta.get("name", ""),
                    "country": meta.get("country", ""),
                    "org": meta.get("org", ""),
                    "dob": meta.get("dob", ""),
                    "rank": meta.get("rank", ""),
                    "region": (meta.get("region") or meta.get("country") or ""),
                    "finish_f": ff,
                    "finish": fmt_ru_time(finish),
                    "status": str(a.get("status") or ""),
                    "splits": sp,
                })
                idx_global += 1

        if self.sort_mode_cb.currentIndex() == 1:
            rows.sort(key=lambda r: (r["finish_f"] is None, r["idx"]))
        else:
            rows.sort(key=lambda r: (r["finish_f"] is None, r["finish_f"] if r["finish_f"] is not None else 0.0, r["idx"]))

        place = 0
        for r in rows:
            st_u = (r.get("status") or "").strip().upper()
            if r.get("finish_f") is not None and st_u != "DNS":
                place += 1
                r["place"] = str(place)
            else:
                r["place"] = ""

        return rows, sorted(split_ids_set, key=split_sort_key)

    def _checkpoint_plan_for_group(self, disc: str, group_runs: List[Tuple[str, Dict[str, Any]]]) -> Tuple[int, Optional[int], List[str]]:
        by_type = {
            "Гит 125 м с места": 125,
            "Гит 250 м с места": 250,
            "Гит 500 м с места": 500,
            "Гит 1000 м с места": 1000,
            "Гонка преследования 2 км": 2000,
            "Гонка преследования 3 км": 3000,
            "Гонка преследования 4 км": 4000,
        }
        dm = by_type.get(disc)
        if dm is None:
            for _k, run in group_runs:
                dm = parse_distance_m(run)
                if dm:
                    break
        if dm:
            cp = max(1, int(round(dm / DIST_STEP_M)))
        else:
            cp = 0
            for _k, run in group_runs:
                cp = max(cp, self._run_checkpoint_count(run))
            cp = max(cp, 1)

        labels: List[str] = []
        # Spreadsheet style close to official protocol samples:
        # 125m -> [125м], 250m -> [250м], 500m -> [250м, 250м], etc.
        if dm and dm >= 250 and (dm % 250 == 0):
            cp = max(1, int(dm // 250))
            labels = ["250м"] * cp
        elif dm == 125:
            cp = 1
            labels = ["125 м"]
        else:
            for i_cp in range(1, cp + 1):
                if i_cp == 1:
                    labels.append(f"{DIST_STEP_M} м")
                else:
                    a = DIST_STEP_M * (i_cp - 1)
                    b = DIST_STEP_M * i_cp
                    labels.append(f"{a}-{b}м")
        return cp, dm, labels

    def _build_all_protocol_text(self, items: Optional[List[Tuple[str, Dict[str, Any]]]] = None) -> str:
        groups = self._protocol_groups(items)
        if not groups:
            return ""
        blocks = [self._protocol_text_for_group(gk, runs).rstrip() for gk, runs in groups]
        return ("\n" + ("-" * 72) + "\n").join(blocks) + "\n"

    def _protocol_text_for_group(self, group_key: Tuple[str, str], group_runs: List[Tuple[str, Dict[str, Any]]]) -> str:
        disc, cat = group_key
        ev = self.event_ed.text().strip() or disc
        rnd = self.round_ed.text().strip()
        sec = self.secretary_ed.text().strip()
        judge = self.judge_ed.text().strip()
        notes = self.notes_ed.text().strip()
        run_keys = [k for k, _r in group_runs]
        rows, split_ids_sorted = self._protocol_group_rows(group_runs)

        lines = ["ПРОТОКОЛ"]
        if ev:
            lines.append(f"Дисциплина: {ev}")
        if cat:
            lines.append(f"Категория: {cat}")
        if rnd:
            lines.append(f"Этап: {rnd}")
        if self.date_ed.text().strip():
            lines.append(self.date_ed.text().strip())
        if self.cond_ed.text().strip():
            lines.append(self.cond_ed.text().strip())
        if run_keys:
            lines.append("Заезды: " + ", ".join(run_keys))
        if judge:
            lines.append(f"Главный судья: {judge}")
        if sec:
            lines.append(f"Секретарь: {sec}")
        lines.append("")

        headers = ["Место", "№", "Имя", "Стр/Гор", "Финиш", "Статус"]
        if self.include_splits_cb.isChecked() and split_ids_sorted:
            headers.append("Отсечки")
        lines.append(";".join(headers))

        for r in rows:
            row = [str(r.get("place") or ""), str(r.get("bib") or ""), str(r.get("name") or ""), str(r.get("country") or ""), str(r.get("finish") or ""), str(r.get("status") or "")]
            if self.include_splits_cb.isChecked() and split_ids_sorted:
                sp_map = r.get("splits") if isinstance(r.get("splits"), dict) else {}
                row.append(" ".join([f"S{sid}:{fmt_ru_time(sp_map.get(str(sid)))}" for sid in split_ids_sorted if str(sid) in sp_map]))
            lines.append(";".join(row))
        if notes:
            lines.append("")
            lines.append(f"Примечание: {notes}")
        lines.append("")
        return "\n".join(lines)

    def _protocol_text_for_run(self, run_key: str, run: Dict[str, Any]) -> str:
        return self._protocol_text_for_group((self._effective_protocol_type_for_run(run_key, run) or infer_discipline(run) or "Произвольно", str(run.get("category") or self.run_categories.get(run_key) or "").strip() or "Без категории"), [(run_key, run)])

    def export_protocol_txt(self):
        items = self._protocol_items()
        if not items:
            QMessageBox.critical(self, "Ошибка", "Нет заездов для экспорта")
            return
        path, _ = QFileDialog.getSaveFileName(self, "Сохранить TXT", "protocol_all.txt", "Text (*.txt)")
        if not path:
            return
        text = self._build_all_protocol_text(items)
        Path(path).write_text(text, encoding="utf-8")
        QMessageBox.information(self, "Готово", f"Сохранено: {path}")

    def _copy_selected_ath_cells(self):
        idxs = self.ath_table.selectedIndexes()
        if not idxs:
            it = self.ath_table.currentItem()
            if not it:
                return
            QApplication.clipboard().setText(str(it.text() or ""))
            return

        rows = sorted({i.row() for i in idxs})
        cols = sorted({i.column() for i in idxs})
        grid = {(i.row(), i.column()): i for i in idxs}
        lines: List[str] = []
        for r in rows:
            vals: List[str] = []
            for c in cols:
                i = grid.get((r, c))
                vals.append(str(i.data() or "") if i is not None else "")
            lines.append("\t".join(vals))
        QApplication.clipboard().setText("\n".join(lines))

    def _copy_protocol_text(self):
        text = self._build_all_protocol_text(self._protocol_items())
        if not text.strip():
            return
        cb = QApplication.clipboard()
        cb.setText(text)
        self.protocol_preview.setPlainText(text)

    def export_protocol_docx(self):
        if Document is None:
            QMessageBox.critical(self, "Ошибка", "python-docx не установлен")
            return
        items = self._protocol_items()
        groups = self._protocol_groups(items)
        if not groups:
            QMessageBox.critical(self, "Ошибка", "Нет заездов для экспорта")
            return
        path, _ = QFileDialog.getSaveFileName(self, "Сохранить DOCX", "protocol_all.docx", "DOCX (*.docx)")
        if not path:
            return
        doc = Document()
        for i, (gk, runs) in enumerate(groups):
            if i > 0:
                doc.add_page_break()
            for ln in self._protocol_text_for_group(gk, runs).splitlines():
                doc.add_paragraph(ln)
        doc.save(path)
        QMessageBox.information(self, "Готово", f"Сохранено: {path}")

    def export_protocol_docx_folder(self):
        if Document is None:
            QMessageBox.critical(self, "Ошибка", "python-docx не установлен")
            return
        items = self._protocol_items()
        groups = self._protocol_groups(items)
        if not groups:
            QMessageBox.critical(self, "Ошибка", "Нет заездов для экспорта")
            return
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения протоколов")
        if not folder:
            return

        def safe_name(s: str) -> str:
            s = (s or "").strip() or "protocol"
            s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
            s = re.sub(r"\s+", " ", s).strip()
            return s[:120]

        saved = 0
        for (disc, cat), runs in groups:
            doc = Document()
            for ln in self._protocol_text_for_group((disc, cat), runs).splitlines():
                doc.add_paragraph(ln)
            base = f"{disc or 'protocol'}_{cat or 'Без категории'}"
            out_path = Path(folder) / f"{safe_name(base)}.docx"
            doc.save(str(out_path))
            saved += 1
        QMessageBox.information(self, "Готово", f"Сохранено файлов: {saved}\nПапка: {folder}")

    def export_protocol_xlsx(self):
        if load_workbook is None:
            QMessageBox.critical(self, "Ошибка", "openpyxl не установлен")
            return
        items = self._protocol_items()
        groups = self._protocol_groups(items)
        if not groups:
            QMessageBox.critical(self, "Ошибка", "Нет заездов для экспорта")
            return
        tpl_path = Path("протокол.xlsx")
        if not tpl_path.exists():
            p, _ = QFileDialog.getOpenFileName(self, "Выбери шаблон", "", "Excel (*.xlsx)")
            if not p:
                return
            tpl_path = Path(p)
        out, _ = QFileDialog.getSaveFileName(self, "Сохранить XLSX", "protocol_all.xlsx", "Excel (*.xlsx)")
        if not out:
            return

        wb = load_workbook(str(tpl_path))
        ws = wb.active
        template_top = 1
        template_bottom = min(25, ws.max_row if ws.max_row > 0 else 25)
        template_height = max(1, template_bottom - template_top + 1)
        template_max_col = max(11, ws.max_column if ws.max_column > 0 else 11)

        plans: List[Dict[str, Any]] = []
        max_checkpoints = 1
        for (disc, cat), group_runs in groups:
            cp_count, dm, labels = self._checkpoint_plan_for_group(disc, group_runs)
            rows, split_ids = self._protocol_group_rows(group_runs)
            plans.append({
                "disc": disc,
                "cat": cat,
                "checkpoints": cp_count,
                "distance_m": dm,
                "labels": labels,
                "rows": rows,
                "split_ids": split_ids,
            })
            max_checkpoints = max(max_checkpoints, int(cp_count or 1))

        block_max_col = 7 + max_checkpoints + 2

        def _copy_block(dst_top: int):
            for r_off in range(template_height):
                src_r = template_top + r_off
                dst_r = dst_top + r_off
                src_dim = ws.row_dimensions.get(src_r)
                dst_dim = ws.row_dimensions[dst_r]
                if src_dim and src_dim.height is not None:
                    dst_dim.height = src_dim.height
                for c in range(1, block_max_col + 1):
                    src_c = c if c <= template_max_col else template_max_col
                    src = ws.cell(src_r, src_c)
                    dst = ws.cell(dst_r, c)
                    dst.value = src.value if c <= template_max_col else None
                    dst.number_format = src.number_format
                    dst.font = copy.copy(src.font)
                    dst.fill = copy.copy(src.fill)
                    dst.border = copy.copy(src.border)
                    dst.alignment = copy.copy(src.alignment)
                    dst.protection = copy.copy(src.protection)

            merges = list(ws.merged_cells.ranges)
            for rng in merges:
                if rng.min_row >= template_top and rng.max_row <= template_bottom:
                    dr = dst_top - template_top
                    shifted = copy.copy(rng)
                    shifted.shift(row_shift=dr, col_shift=0)
                    try:
                        ws.merge_cells(str(shifted))
                    except Exception:
                        pass

        needed_rows = template_height * len(groups)
        if ws.max_row < needed_rows:
            ws.insert_rows(ws.max_row + 1, amount=(needed_rows - ws.max_row))

        for i_blk in range(1, len(groups)):
            _copy_block(1 + i_blk * template_height)

        for i_group, plan in enumerate(plans):
            block_top = 1 + i_group * template_height
            disc = str(plan.get("disc") or "Произвольно")
            cat = str(plan.get("cat") or "")
            cp_count = int(plan.get("checkpoints") or 1)
            labels = list(plan.get("labels") or [])
            dm = plan.get("distance_m")
            rows = list(plan.get("rows") or [])
            split_ids = list(plan.get("split_ids") or [])

            title = (self.event_ed.text().strip() or disc).upper()
            stage = (self.round_ed.text().strip() or "1 ЭТАП").upper()

            ws.cell(row=block_top + 0, column=1, value=title)
            ws.cell(row=block_top + 1, column=1, value=stage)
            ws.cell(row=block_top + 2, column=1, value="РЕЗУЛЬТАТЫ")
            ws.cell(row=block_top + 3, column=1, value=self.date_ed.text().strip())
            ws.cell(row=block_top + 4, column=1, value=self.cond_ed.text().strip())

            hdr_row = block_top + 5
            base_headers = ["Ме\nсто", "№\nг-ка", "Фамилия  Имя Отчество", "Организация", "Дата\nРождения", "Раз\nряд", "Регион"]
            for j, h in enumerate(base_headers, start=1):
                ws.cell(row=hdr_row, column=j, value=h)

            split_col_start = 8
            split_col_end = split_col_start + max_checkpoints - 1
            for j in range(split_col_start, split_col_end + 1):
                ws.cell(row=hdr_row, column=j, value=None)
            for j, lab in enumerate(labels, start=split_col_start):
                ws.cell(row=hdr_row, column=j, value=lab)

            result_col = split_col_start + cp_count
            speed_col = result_col + 1
            ws.cell(row=hdr_row, column=result_col, value="Результ\nтат")
            ws.cell(row=hdr_row, column=speed_col, value="Ср.ск-ть")

            data_start = block_top + 7
            data_end = block_top + 24
            for rr_i in range(data_start, data_end + 1):
                for c in range(1, block_max_col + 1):
                    ws.cell(row=rr_i, column=c, value=None)

            for i_row, rr in enumerate(rows):
                row_idx = data_start + i_row
                if row_idx > data_end:
                    break
                ff = rr.get("finish_f")
                splits = rr.get("splits") if isinstance(rr.get("splits"), dict) else {}
                cumulative: List[Optional[float]] = []
                for i_cp in range(1, cp_count + 1):
                    if i_cp == cp_count:
                        cumulative.append(ff)
                    else:
                        val = None
                        if i_cp - 1 < len(split_ids):
                            try:
                                sv = splits.get(str(split_ids[i_cp - 1]))
                                val = float(sv) if sv is not None else None
                            except Exception:
                                val = None
                        cumulative.append(val)
                prev = 0.0
                cp_vals: List[str] = []
                for i_cp, cur in enumerate(cumulative):
                    seg = None
                    if cur is not None:
                        seg = cur if i_cp == 0 else (cur - prev if prev is not None else None)
                        prev = cur
                    else:
                        prev = None
                    cp_vals.append(fmt_sec_ru(seg) if seg is not None else "")

                vals = [
                    str(rr.get("place") or ""),
                    str(rr.get("bib") or ""),
                    str(rr.get("name") or ""),
                    str(rr.get("org") or ""),
                    str(rr.get("dob") or ""),
                    str(rr.get("rank") or ""),
                    str(rr.get("region") or rr.get("country") or ""),
                ]
                vals += cp_vals
                vals += [fmt_sec_ru(ff) if ff is not None else str(rr.get("status") or ""), fmt_speed_kmh_ru(dm, ff) if ff is not None else ""]
                for j, v in enumerate(vals, start=1):
                    ws.cell(row=row_idx, column=j, value=v)

        wb.save(out)
        QMessageBox.information(self, "Готово", f"Сохранено: {out}")

    def _pump(self):
        try:
            while True:
                kind, payload = self.q.get_nowait()
                if kind == "status" and payload == "connected":
                    self.scan_stop_evt.set()
                    self._set_status("Подключено", True)
                elif kind == "state" and isinstance(payload, dict):
                    self.state = payload
                    runs = payload.get("runs") if isinstance(payload.get("runs"), dict) else {}
                    self.run_categories = {}
                    for rk, rv in runs.items():
                        if isinstance(rv, dict):
                            cat = str(rv.get("category") or "").strip()
                            if cat:
                                self.run_categories[str(rk)] = cat
                    self._refresh_categories()
                    if not self.selected_run_key:
                        self.selected_run_key = payload.get("current_key")
                    self._refresh_views()
                elif kind == "err":
                    self._set_status("Ошибка", False)
                    self._append_log("ERROR: " + str(payload))
                elif kind == "scan_status":
                    self._set_status(str(payload), False)
                    self._append_log(str(payload))
                elif kind == "scan_found":
                    host, port = payload
                    self.host_ed.setText(str(host))
                    self.port_ed.setText(str(port))
                    self._append_log(f"SCAN: найден сервер {host}:{port}")
                    self.connect_net(silent=True)
                elif kind == "scan_none":
                    self._append_log("SCAN: сервер не найден")
        except queue.Empty:
            pass

        now = time.monotonic()
        net_alive = bool(self.net and self.net.is_alive())
        scan_alive = bool(self.scan_thr and self.scan_thr.is_alive())
        if self.auto_connect_cb.isChecked() and (not self._connected) and (not net_alive) and (not scan_alive):
            if (now - self._last_connect_try) >= 3.0:
                if self.auto_scan_cb.isChecked():
                    self.start_scan_connect()
                else:
                    self.connect_net(silent=True)

    def closeEvent(self, event):
        self._save_settings()
        try:
            self._save_roster(silent=True)
        except Exception:
            pass
        self.scan_stop_evt.set()
        self.disconnect_net()
        super().closeEvent(event)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--host", default="127.0.0.1")
    ap.add_argument("--port", type=int, default=9876)
    args = ap.parse_args()

    app = QApplication([])
    w = ClientQtApp(args.host, args.port)
    w.show()
    app.exec()


if __name__ == "__main__":
    main()
